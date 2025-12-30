"""
DOCX 格式感知分片实现示例

本文件展示了关键实现代码，供参考使用。
实际实现时，请将这些代码集成到 collabtrans/translator/ai_translator/docx_translator.py 中。
"""

from dataclasses import dataclass
from typing import List, Dict, Any, Tuple
from docx.text.run import Run
from docx.text.paragraph import Paragraph
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
import logging

logger = logging.getLogger(__name__)


@dataclass
class RunFormatInfo:
    """Run 格式信息"""
    font_name: str | None = None
    font_size: int | None = None  # 单位：磅 (Pt)
    bold: bool | None = None
    italic: bool | None = None
    underline: bool | None = None
    color_rgb: str | None = None  # RGB 颜色值，格式："RRGGBB" 或 "RGBColor(r, g, b)"
    highlight_color: str | None = None  # 高亮颜色
    strikethrough: bool | None = None
    
    def __eq__(self, other) -> bool:
        """比较两个格式信息是否相同（用于分片判断）"""
        if not isinstance(other, RunFormatInfo):
            return False
        return (
            self.font_name == other.font_name and
            self.font_size == other.font_size and
            self.bold == other.bold and
            self.italic == other.italic and
            self.underline == other.underline and
            self.color_rgb == other.color_rgb and
            self.highlight_color == other.highlight_color and
            self.strikethrough == other.strikethrough
        )
    
    def __hash__(self) -> int:
        """用于字典键"""
        return hash((
            self.font_name, self.font_size, self.bold, self.italic,
            self.underline, self.color_rgb, self.highlight_color, self.strikethrough
        ))


def extract_run_format(run: Run) -> RunFormatInfo:
    """从 Run 对象提取格式信息"""
    font = run.font
    
    # 提取字体名称
    font_name = None
    if font.name:
        font_name = font.name
    
    # 提取字号
    font_size = None
    if font.size:
        font_size = font.size.pt
    
    # 提取颜色（RGB）
    color_rgb = None
    if font.color and font.color.rgb:
        # 转换为字符串表示
        rgb = font.color.rgb
        if isinstance(rgb, RGBColor):
            color_rgb = f"RGBColor({rgb.r}, {rgb.g}, {rgb.b})"
        else:
            color_rgb = str(rgb)
    
    # 提取高亮颜色
    highlight_color = None
    if hasattr(font, 'highlight_color') and font.highlight_color:
        highlight_color = str(font.highlight_color)
    
    # 提取删除线
    strikethrough = None
    if hasattr(font, 'strike'):
        strikethrough = font.strike
    
    return RunFormatInfo(
        font_name=font_name,
        font_size=font_size,
        bold=font.bold,
        italic=font.italic,
        underline=font.underline,
        color_rgb=color_rgb,
        highlight_color=highlight_color,
        strikethrough=strikethrough
    )


def process_paragraph_format_aware(para: Paragraph, is_image_run_func) -> Tuple[List[Dict[str, Any]], List[str]]:
    """
    格式感知的段落处理
    
    Args:
        para: 段落对象
        is_image_run_func: 判断 run 是否为图片的函数
    
    Returns:
        (elements_to_translate, original_texts) 元组
    """
    elements_to_translate = []
    original_texts = []
    
    current_text_segment = ""
    current_runs = []
    current_format: RunFormatInfo | None = None
    
    for run in para.runs:
        if is_image_run_func(run):
            # 遇到图片，结束当前分片
            if current_text_segment.strip():
                elements_to_translate.append({
                    "type": "text_runs",
                    "runs": current_runs,
                    "format": current_format  # 记录格式信息
                })
                original_texts.append(current_text_segment)
            # 重置
            current_text_segment = ""
            current_runs = []
            current_format = None
        else:
            # 提取当前 run 的格式
            run_format = extract_run_format(run)
            
            # 判断是否需要分片
            # 条件：格式发生变化 或 当前分片为空
            if current_format is None:
                # 第一个 run，开始新分片
                current_format = run_format
                current_runs.append(run)
                current_text_segment += run.text
            elif current_format == run_format:
                # 格式相同，继续累积
                current_runs.append(run)
                current_text_segment += run.text
            else:
                # 格式不同，结束当前分片，开始新分片
                if current_text_segment.strip():
                    elements_to_translate.append({
                        "type": "text_runs",
                        "runs": current_runs,
                        "format": current_format
                    })
                    original_texts.append(current_text_segment)
                
                # 开始新分片
                current_format = run_format
                current_runs = [run]
                current_text_segment = run.text
    
    # 处理最后一个分片
    if current_text_segment.strip():
        elements_to_translate.append({
            "type": "text_runs",
            "runs": current_runs,
            "format": current_format
        })
        original_texts.append(current_text_segment)
    
    return elements_to_translate, original_texts


def apply_format_to_run(run: Run, format_info: RunFormatInfo, target_font_name: str | None = None):
    """
    将格式信息应用到 Run 对象
    
    Args:
        run: Run 对象
        format_info: 格式信息
        target_font_name: 目标语言字体名称（可选，用于字体兼容性）
    """
    font = run.font
    
    # 应用字体名称（优先使用目标语言字体，如果未指定则使用原始字体）
    if target_font_name:
        font.name = target_font_name
    elif format_info.font_name:
        font.name = format_info.font_name
    
    # 应用字号
    if format_info.font_size:
        font.size = Pt(format_info.font_size)
    
    # 应用粗体
    if format_info.bold is not None:
        font.bold = format_info.bold
    
    # 应用斜体
    if format_info.italic is not None:
        font.italic = format_info.italic
    
    # 应用下划线
    if format_info.underline is not None:
        font.underline = format_info.underline
    
    # 应用颜色
    if format_info.color_rgb:
        try:
            # 解析 RGB 颜色字符串
            if format_info.color_rgb.startswith("RGBColor"):
                # 从 "RGBColor(r, g, b)" 提取
                import re
                match = re.search(r'RGBColor\((\d+),\s*(\d+),\s*(\d+)\)', format_info.color_rgb)
                if match:
                    r, g, b = map(int, match.groups())
                    font.color.rgb = RGBColor(r, g, b)
            else:
                # 从 "RRGGBB" 格式提取（十六进制）
                if len(format_info.color_rgb) == 6:
                    r = int(format_info.color_rgb[0:2], 16)
                    g = int(format_info.color_rgb[2:4], 16)
                    b = int(format_info.color_rgb[4:6], 16)
                    font.color.rgb = RGBColor(r, g, b)
        except Exception as e:
            logger.warning(f"Failed to apply color {format_info.color_rgb}: {e}")
    
    # 应用高亮颜色
    if format_info.highlight_color:
        try:
            # 注意：python-docx 对高亮颜色的支持可能有限
            # 这里需要根据实际情况调整
            if hasattr(font, 'highlight_color'):
                # 尝试解析高亮颜色
                pass
        except Exception as e:
            logger.warning(f"Failed to apply highlight color: {e}")
    
    # 应用删除线
    if format_info.strikethrough is not None:
        if hasattr(font, 'strike'):
            font.strike = format_info.strikethrough


def preserve_page_breaks_in_run(run: Run, new_text: str) -> None:
    """设置 run 文本，同时保留分页符"""
    # 检查 run 是否有分页符
    from docx.oxml.ns import qn
    breaks = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')
    page_breaks = [
        br for br in breaks 
        if br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page'
    ]
    
    if page_breaks:
        # 清除现有文本元素但保留分页符
        text_elements = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
        for text_elem in text_elements:
            text_elem.text = new_text
            # 只保留第一个文本元素，移除其他
            break
        else:
            # 未找到文本元素，创建一个
            from docx.oxml import OxmlElement
            t_elem = OxmlElement('w:t')
            t_elem.text = new_text
            run._element.append(t_elem)
    else:
        # 没有分页符，直接设置文本
        run.text = new_text


# ===== 在 DocxTranslator 类中的使用示例 =====

class DocxTranslator:
    """示例：如何在 DocxTranslator 中集成格式感知分片"""
    
    def _pre_translate(self, document) -> Tuple:
        """
        修改后的 _pre_translate 方法
        使用格式感知的分片逻辑
        """
        from io import BytesIO
        import docx
        from docx.document import Document as DocumentObject
        
        doc = docx.Document(BytesIO(document.content))
        elements_to_translate = []
        original_texts = []
        skipped_toc = 0
        
        def process_paragraph(para: Paragraph):
            nonlocal elements_to_translate, original_texts, skipped_toc
            
            # 跳过 TOC 段落（保持原有逻辑）
            if self._paragraph_has_toc_field(para):
                skipped_toc += 1
                return
            
            # 使用格式感知的分片逻辑
            para_elements, para_texts = process_paragraph_format_aware(
                para, 
                is_image_run_func=self.is_image_run
            )
            elements_to_translate.extend(para_elements)
            original_texts.extend(para_texts)
        
        # 遍历所有段落
        for para in doc.paragraphs:
            process_paragraph(para)
        
        # 遍历所有表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        process_paragraph(para)
        
        return doc, elements_to_translate, original_texts
    
    def _after_translate(
        self, 
        doc, 
        elements_to_translate: List[Dict[str, Any]],
        translated_texts: List[str], 
        original_texts: List[str]
    ) -> bytes:
        """
        修改后的 _after_translate 方法
        使用格式感知的导出逻辑
        """
        from io import BytesIO
        from collabtrans.translator.ai_translator.docx_translator import get_font_for_language
        
        translation_map = dict(zip(original_texts, translated_texts))
        
        # 获取目标语言字体（用于字体兼容性）
        target_font_name = get_font_for_language(self.config.to_lang)
        
        for i, element_info in enumerate(elements_to_translate):
            runs = element_info["runs"]
            format_info: RunFormatInfo = element_info.get("format")
            original_text = original_texts[i]
            translated_text = translated_texts[i]
            
            # 确定最终文本（根据 insert_mode）
            if self.insert_mode == "replace":
                final_text = translated_text
            elif self.insert_mode == "append":
                final_text = original_text + self.separator + translated_text
            elif self.insert_mode == "prepend":
                final_text = translated_text + self.separator + original_text
            else:
                final_text = translated_text
            
            if not runs:
                continue
            
            # 策略：将翻译文本分配到各个 run，保持格式
            if len(runs) == 1:
                # 单个 run：直接替换并应用格式
                first_run = runs[0]
                preserve_page_breaks_in_run(first_run, final_text)
                if format_info:
                    apply_format_to_run(first_run, format_info, target_font_name)
            else:
                # 多个 run：将文本放入第一个 run，其他 run 清空
                first_run = runs[0]
                preserve_page_breaks_in_run(first_run, final_text)
                if format_info:
                    apply_format_to_run(first_run, format_info, target_font_name)
                
                # 清空其他 run（保留结构）
                for run in runs[1:]:
                    run.text = ""
                    # 可选：也应用格式（保持一致性）
                    if format_info:
                        apply_format_to_run(run, format_info, target_font_name)
        
        # 保存文档
        doc_output_stream = BytesIO()
        doc.save(doc_output_stream)
        return doc_output_stream.getvalue()
