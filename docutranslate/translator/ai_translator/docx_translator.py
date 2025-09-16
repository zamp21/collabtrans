# SPDX-FileCopyrightText: 2025 QinHan
# SPDX-License-Identifier: MPL-2.0
import asyncio
from dataclasses import dataclass
from io import BytesIO
from typing import Self, Literal, List, Dict, Any, Tuple

import docx
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from docutranslate.agents.segments_agent import SegmentsTranslateAgentConfig, SegmentsTranslateAgent
from docutranslate.ir.document import Document
from docutranslate.translator.ai_translator.base import AiTranslatorConfig, AiTranslator


def is_image_run(run: Run) -> bool:
    """检查一个 run 是否包含图片。"""
    # w:drawing 是嵌入式图片的标志, w:pict 是 VML 图片的标志
    return '<w:drawing' in run.element.xml or '<w:pict' in run.element.xml


@dataclass
class DocxTranslatorConfig(AiTranslatorConfig):
    """
    DocxTranslator 的配置类。
    """
    insert_mode: Literal["replace", "append", "prepend"] = "replace"
    separator: str = "\n"


class DocxTranslator(AiTranslator):
    """
    用于翻译 .docx 文件的翻译器。
    此版本经过优化，可以处理图文混排的段落而不会丢失图片。
    """

    def __init__(self, config: DocxTranslatorConfig):
        super().__init__(config=config)
        self.chunk_size = config.chunk_size
        self.translate_agent = None
        if not self.skip_translate:
            agent_config = SegmentsTranslateAgentConfig(
                custom_prompt=config.custom_prompt,
                to_lang=config.to_lang,
                base_url=config.base_url,
                api_key=config.api_key,
                model_id=config.model_id,
                temperature=config.temperature,
                thinking=config.thinking,
                concurrent=config.concurrent,
                timeout=config.timeout,
                logger=self.logger,
                glossary_dict=config.glossary_dict,
                retry=config.retry
            )
            self.translate_agent = SegmentsTranslateAgent(agent_config)
        self.insert_mode = config.insert_mode
        self.separator = config.separator

    def _pre_translate(self, document: Document) -> Tuple[DocumentObject, List[Dict[str, Any]], List[str]]:
        """
        [已重构] 预处理 .docx 文件，在 Run 级别上提取文本，以避免破坏图片。
        :param document: 包含 .docx 文件内容的 Document 对象。
        :return: 一个元组，包含：
                 - docx.Document 对象
                 - 一个包含文本块信息的列表 (每个元素代表一组连续的文本 run)
                 - 一个包含所有待翻译原文的列表
        """
        doc = docx.Document(BytesIO(document.content))
        elements_to_translate = []
        original_texts = []

        def process_paragraph(para: Paragraph):
            nonlocal elements_to_translate, original_texts
            current_text_segment = ""
            current_runs = []

            for run in para.runs:
                if is_image_run(run):
                    # 遇到图片，将之前累积的文本作为一个翻译单元
                    if current_text_segment.strip():
                        elements_to_translate.append({"type": "text_runs", "runs": current_runs})
                        original_texts.append(current_text_segment)
                    # 重置累加器
                    current_text_segment = ""
                    current_runs = []
                else:
                    # 累积文本 run
                    current_runs.append(run)
                    current_text_segment += run.text

            # 处理段落末尾的最后一个文本块
            if current_text_segment.strip():
                elements_to_translate.append({"type": "text_runs", "runs": current_runs})
                original_texts.append(current_text_segment)

        # 遍历所有段落
        for para in doc.paragraphs:
            process_paragraph(para)

        # 遍历所有表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        process_paragraph(para)

        return doc, elements_to_translate, original_texts

    def _after_translate(self, doc: DocumentObject, elements_to_translate: List[Dict[str, Any]],
                         translated_texts: List[str], original_texts: List[str]) -> bytes:
        """
        [已重构] 将翻译后的文本写回到对应的 text runs 中，保留图片和样式。
        """
        translation_map = dict(zip(original_texts, translated_texts))

        for i, element_info in enumerate(elements_to_translate):
            runs = element_info["runs"]
            original_text = original_texts[i]
            translated_text = translated_texts[i]

            # 根据插入模式确定最终文本
            if self.insert_mode == "replace":
                final_text = translated_text
            elif self.insert_mode == "append":
                final_text = original_text + self.separator + translated_text
            elif self.insert_mode == "prepend":
                final_text = translated_text + self.separator + original_text
            else:
                self.logger.error("不正确的DocxTranslatorConfig参数")
                final_text = translated_text

            if not runs:
                continue

            # --- 这是修改的核心部分 ---
            # 1. 将完整的翻译文本写入第一个 run
            first_run = runs[0]
            first_run.text = final_text

            # 2. 清空该文本块中其余 run 的内容，但保留 run 本身及其格式
            #    这可以防止重复文本，同时保留文档结构
            for run in runs[1:]:
                run.text = ""
            # --- 修改结束 ---

        # 将修改后的文档保存到 BytesIO 流
        doc_output_stream = BytesIO()
        doc.save(doc_output_stream)
        return doc_output_stream.getvalue()

    def translate(self, document: Document) -> Self:
        """
        同步翻译 .docx 文件。
        """
        doc, elements_to_translate, original_texts = self._pre_translate(document)
        if not original_texts:
            print("\n文件中没有找到需要翻译的文本内容。")
            output_stream = BytesIO()
            doc.save(output_stream)
            document.content = output_stream.getvalue()
            return self

        if self.glossary_agent:
            self.glossary_dict_gen = self.glossary_agent.send_segments(original_texts, self.chunk_size)
            if self.translate_agent:
                self.translate_agent.update_glossary_dict(self.glossary_dict_gen)

        # 调用翻译 agent
        if self.translate_agent:
            translated_texts = self.translate_agent.send_segments(original_texts, self.chunk_size)
        else:
            translated_texts = original_texts

        # 将翻译结果写回文档
        document.content = self._after_translate(doc, elements_to_translate, translated_texts, original_texts)
        return self

    async def translate_async(self, document: Document) -> Self:
        """
        异步翻译 .docx 文件。
        """
        doc, elements_to_translate, original_texts = await asyncio.to_thread(self._pre_translate, document)
        if not original_texts:
            print("\n文件中没有找到需要翻译的文本内容。")
            # 在异步环境中正确保存和返回
            output_stream = BytesIO()
            doc.save(output_stream)
            document.content = output_stream.getvalue()
            return self

        if self.glossary_agent:
            self.glossary_dict_gen = await self.glossary_agent.send_segments_async(original_texts, self.chunk_size)
            if self.translate_agent:
                self.translate_agent.update_glossary_dict(self.glossary_dict_gen)

        # 异步调用翻译 agent
        if self.translate_agent:
            translated_texts = await self.translate_agent.send_segments_async(original_texts, self.chunk_size)
        else:
            translated_texts = original_texts
        # 将翻译结果写回文档
        document.content = await asyncio.to_thread(self._after_translate, doc, elements_to_translate, translated_texts,
                                                   original_texts)
        return self
