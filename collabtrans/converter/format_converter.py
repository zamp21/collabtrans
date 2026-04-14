"""
Document format converter module.
Supports PDF to DOCX conversion using MinerU + Pandoc.
"""

import asyncio
import logging
import os
import re
import tempfile
import uuid
import subprocess
import time
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime, timedelta

logger = logging.getLogger(__name__)

# Global storage for conversion tasks
conversion_tasks: Dict[str, Dict[str, Any]] = {}
conversion_files: Dict[str, Dict[str, Any]] = {}  # Store file info for cleanup


class ConversionError(Exception):
    """Custom exception for conversion errors"""
    pass


class FormatConverter:
    """Document format converter"""
    
    def __init__(self):
        self.supported_formats = {
            'pdf': ['docx'],
            # Future: 'docx': ['pdf'], 'txt': ['docx']
        }
    
    async def _send_log(self, log_queue: Optional[asyncio.Queue], message: str):
        """Send log message to frontend if log_queue is provided"""
        if log_queue:
            try:
                await log_queue.put(message)
            except Exception as e:
                logger.warning(f"Failed to send log to frontend: {e}")
    
    def get_supported_targets(self, source_format: str) -> list[str]:
        """Get supported target formats for a source format"""
        return self.supported_formats.get(source_format.lower(), [])
    
    async def convert_pdf_to_docx(
        self,
        pdf_path: str,
        output_path: str,
        quality: str = 'high',
        log_queue: Optional[asyncio.Queue] = None,
        use_mineru: bool = False,
        mineru_config: Optional[dict] = None
    ) -> None:
        """Convert PDF to DOCX format with optimized settings

        Args:
            pdf_path: Path to input PDF file
            output_path: Path to output DOCX file
            quality: Conversion quality (not used, kept for compatibility)
            log_queue: Optional queue for sending log messages to frontend
            use_mineru: Whether to use MinerU for OCR
            mineru_config: MinerU configuration dict with keys:
                - api_url: MinerU API URL (e.g., 'http://localhost:8920')
                - model_version: Model version ('vlm' or 'pipeline')
                - formula_ocr: Enable formula OCR
                - ocr_enabled: Enable OCR
                - mineru_token: API token (optional for local deployment)
        """
        import os

        try:
            logger.info(f"Starting PDF to DOCX conversion: {pdf_path} -> {output_path}")
            await self._send_log(log_queue, "Starting PDF to DOCX conversion...")

            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            if use_mineru:
                # Use MinerU + Pandoc for conversion
                await self._send_log(log_queue, "Using MinerU + Pandoc for conversion...")
                await self._convert_pdf_to_docx_with_mineru(pdf_path, output_path, log_queue, mineru_config)
            else:
                # Use traditional pdf2docx conversion
                await self._convert_pdf_to_docx_with_pdf2docx(pdf_path, output_path, log_queue)

            logger.info(f"PDF to DOCX conversion completed: {output_path}")
            await self._send_log(log_queue, "Conversion completed!")

        except Exception as e:
            logger.error(f"PDF to DOCX conversion failed: {e}")
            raise ConversionError(f"Conversion failed: {str(e)}")

    async def _convert_pdf_to_docx_with_mineru(
        self,
        pdf_path: str,
        output_path: str,
        log_queue: Optional[asyncio.Queue] = None,
        mineru_config: Optional[dict] = None
    ) -> None:
        """Convert PDF to DOCX using MinerU (PDF→Markdown/ContentList) + python-docx"""

        # Load MinerU configuration from global_config if not provided
        if not mineru_config:
            try:
                from collabtrans.config.global_config import get_global_config
                gc = get_global_config()
                mineru_engine = gc.translator_settings.engines.get('mineru', {})
                mineru_config = {
                    'api_url': mineru_engine.get('api_url', 'http://localhost:8920'),
                    'model_version': mineru_engine.get('model_version', gc.translator_settings.mineru_model_version),
                    'formula_ocr': gc.translator_settings.formula_ocr,
                    'ocr_enabled': True,
                    'mineru_token': gc.translator_mineru_token
                }
            except Exception as e:
                logger.warning(f"Failed to load MinerU config from global_config: {e}")
                mineru_config = {
                    'api_url': 'http://localhost:8920',
                    'model_version': 'vlm',
                    'formula_ocr': True,
                    'ocr_enabled': True,
                    'mineru_token': ''
                }

        # Get MinerU configuration
        api_url = mineru_config.get('api_url', 'http://localhost:8920')
        model_version = mineru_config.get('model_version', 'vlm')
        formula_ocr = mineru_config.get('formula_ocr', True)
        ocr_enabled = mineru_config.get('ocr_enabled', True)
        mineru_token = mineru_config.get('mineru_token', '')

        await self._send_log(log_queue, f"MinerU API: {api_url}, Model: {model_version}")

        # Import MinerU converter
        try:
            from collabtrans.converter.x2md.converter_mineru import ConverterMineru, ConverterMineruConfig
        except ImportError as e:
            raise ConversionError(f"MinerU converter not available: {str(e)}")

        # Read PDF document
        from collabtrans.ir.document import Document
        document = Document.from_path(pdf_path)

        # Use cacher with lock to prevent duplicate MinerU calls
        from collabtrans.cacher import md_based_convert_cacher
        config = ConverterMineruConfig(
            mineru_token=mineru_token,
            base_url=api_url,
            model_version=model_version,
            formula_ocr=formula_ocr,
            ocr_enabled=ocr_enabled
        )

        # Create converter - always create fresh converter for Convert workflow
        # to avoid cache pollution from translation workflow
        converter = ConverterMineru(config)

        async def do_convert():
            await self._send_log(log_queue, "Parsing PDF with MinerU...")
            # Run synchronous convert in thread
            result = await asyncio.to_thread(converter.convert, document)
            await self._send_log(log_queue, "PDF parsed successfully")
            return result

        # For Convert workflow: directly call MinerU without using cache
        # This ensures we get clean markdown with base64 images,
        # not placeholder-contaminated markdown from translation workflow
        await self._send_log(log_queue, "Converting PDF with MinerU (bypassing cache)...")
        markdown_doc = await do_convert()

        # Save intermediate markdown to temp directory for debugging
        temp_dir = tempfile.gettempdir()
        base_name = os.path.basename(output_path).replace('.docx', '')
        temp_md_path = os.path.join(temp_dir, f"{base_name}_convert.md")
        try:
            md_content = markdown_doc.content.decode('utf-8')
            with open(temp_md_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            await self._send_log(log_queue, f"Intermediate markdown saved to: {temp_md_path}")
            logger.info(f"Intermediate markdown saved to: {temp_md_path}")
        except Exception as e:
            logger.warning(f"Failed to save intermediate markdown: {e}")

        # Check if result has placeholders (should not happen with fresh MinerU call)
        md_content = markdown_doc.content.decode('utf-8')
        has_placeholders = bool(re.search(r'<ph-[a-zA-Z0-9]+>', md_content))
        if has_placeholders:
            await self._send_log(log_queue, "Warning: MinerU result contains unexpected placeholders")
            logger.warning("Fresh MinerU result contains placeholders - unexpected!")

        # Use Pandoc to convert markdown (handles HTML tables with rowspan/colspan and images properly)
        # Pandoc path is preferred because:
        # 1. HTML tables in markdown are preserved and converted correctly
        # 2. Base64 images are handled correctly
        # 3. python-docx doesn't support rowspan/colspan
        await self._send_log(log_queue, "Converting markdown to DOCX with Pandoc (handles HTML tables and images)...")
        await self._build_docx_from_markdown(markdown_doc, output_path, log_queue)

        await self._send_log(log_queue, "DOCX file created successfully")

    async def _build_docx_from_markdown(
        self,
        markdown_doc,
        output_path: str,
        log_queue: Optional[asyncio.Queue] = None
    ) -> None:
        """Build DOCX from Markdown using Pandoc (fallback method)"""
        import tempfile

        # Save markdown to a temp directory instead of deleting
        temp_dir = tempfile.gettempdir()
        base_name = os.path.basename(output_path).replace('.docx', '')
        temp_md_path = os.path.join(temp_dir, f"{base_name}_convert.md")

        try:
            # Write markdown content to temp file
            md_content = markdown_doc.content.decode('utf-8')
            with open(temp_md_path, 'w', encoding='utf-8') as f:
                f.write(md_content)

            await self._send_log(log_queue, f"Intermediate markdown saved to: {temp_md_path}")
            logger.info(f"Intermediate markdown saved to: {temp_md_path}")

            # Convert markdown to DOCX using two-step process for proper HTML table handling
            # Step 1: MD → HTML (preserves HTML tables with rowspan/colspan)
            # Step 2: HTML → DOCX (converts HTML tables to proper DOCX tables)
            html_path = os.path.join(temp_dir, f"{base_name}_convert.html")

            await self._send_log(log_queue, "Converting markdown to DOCX via HTML (two-step process)")
            logger.info(f"Running pandoc MD→HTML: {temp_md_path} → {html_path}")

            # Step 1: Convert markdown to HTML
            md_to_html_result = subprocess.run(
                [
                    'pandoc',
                    temp_md_path,
                    '-o', html_path,
                    '--from=markdown+pipe_tables+grid_tables+multiline_tables+raw_html+tex_math_dollars',
                    '--to=html',
                ],
                capture_output=True,
                text=True,
                timeout=60
            )

            if md_to_html_result.returncode != 0:
                logger.error(f"Pandoc MD→HTML failed: {md_to_html_result.stderr}")
                raise ConversionError(f"Pandoc MD→HTML conversion failed: {md_to_html_result.stderr}")

            logger.info(f"Pandoc MD→HTML completed successfully")

            # Read and modify HTML to add border attributes for tables
            with open(html_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            # Add border="1" to all table tags (Pandoc HTML→DOCX needs this for visible borders)
            # Properly handle both <table> and <table attrs> cases
            def add_border_attr(match):
                attrs = match.group(1)  # Attributes after 'table', before '>'
                if 'border=' in attrs.lower():
                    return match.group(0)  # Already has border, keep original
                # Add border="1" attribute, keep existing attributes and the closing >
                return '<table border="1"' + attrs + '>'

            html_content = re.sub(
                r'<table\b([^>]*)>',
                add_border_attr,
                html_content
            )

            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)

            await self._send_log(log_queue, "Added table borders to HTML")
            logger.info("Added table border attributes to HTML")

            # Step 2: Convert HTML to DOCX
            logger.info(f"Running pandoc HTML→DOCX: {html_path} → {output_path}")
            html_to_docx_result = subprocess.run(
                [
                    'pandoc',
                    html_path,
                    '-o', output_path,
                    '--from=html',
                    '--to=docx',
                    '--wrap=none',
                ],
                capture_output=True,
                text=True,
                timeout=120
            )

            if html_to_docx_result.returncode != 0:
                logger.error(f"Pandoc HTML→DOCX failed: {html_to_docx_result.stderr}")
                raise ConversionError(f"Pandoc HTML→DOCX conversion failed: {html_to_docx_result.stderr}")

            logger.info(f"Pandoc HTML→DOCX completed successfully")

            await self._send_log(log_queue, f"Pandoc converted HTML→DOCX: {output_path}")

            # Add table borders using python-docx
            logger.info("Starting table border processing with python-docx")
            try:
                from docx import Document as DocxDocument
                from docx.oxml import parse_xml
                from docx.oxml.ns import qn

                logger.info("python-docx imported successfully, loading DOCX file")
                doc = DocxDocument(output_path)
                table_count = len(doc.tables)
                logger.info(f"Found {table_count} tables in DOCX for border processing")
                await self._send_log(log_queue, f"Found {table_count} tables in DOCX")

                for table in doc.tables:
                    tbl = table._tbl
                    tblPr = tbl.tblPr

                    if tblPr is None:
                        tblPr = parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                        tbl.insert(0, tblPr)

                    # Remove tblStyle element to prevent style borders from overriding our borders
                    existing_style = tblPr.find(qn('w:tblStyle'))
                    if existing_style is not None:
                        tblPr.remove(existing_style)
                        logger.debug("Removed table style reference to enable explicit borders")

                    # Create tblBorders element with visible borders
                    tblBorders_xml = parse_xml(
                        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        r'<w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                        r'<w:left w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                        r'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                        r'<w:right w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                        r'<w:insideH w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                        r'<w:insideV w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                        r'</w:tblBorders>'
                    )

                    # Remove existing borders if present
                    existing_borders = tblPr.find(qn('w:tblBorders'))
                    if existing_borders is not None:
                        tblPr.remove(existing_borders)

                    tblPr.append(tblBorders_xml)

                if table_count > 0:
                    doc.save(output_path)
                    await self._send_log(log_queue, f"Added borders to {table_count} tables")
                    logger.info(f"Added table borders to DOCX (found {table_count} tables)")

            except ImportError as e:
                await self._send_log(log_queue, f"Warning: python-docx not available: {e}")
                logger.warning(f"python-docx not available: {e}")
            except Exception as e:
                await self._send_log(log_queue, f"Warning: Failed to add borders: {e}")
                logger.warning(f"Failed to add table borders: {e}")

        except Exception as e:
            logger.error(f"_build_docx_from_markdown error: {e}")
            raise

        # Note: We keep the temp_md_path file for debugging, don't delete it

    async def _convert_pdf_to_docx_with_pdf2docx(
        self,
        pdf_path: str,
        output_path: str,
        log_queue: Optional[asyncio.Queue] = None
    ) -> None:
        """Convert PDF to DOCX using pdf2docx library"""
        import os

        try:
            from pdf2docx import Converter
        except ImportError:
            raise ConversionError("pdf2docx library not installed. Please install it with: pip install pdf2docx")

        cv = Converter(pdf_path)
        await self._send_log(log_queue, "Analyzing PDF document...")

        try:
            if hasattr(cv, 'set_optimization_level'):
                cv.set_optimization_level(1)
        except:
            pass

        await self._send_log(log_queue, "Converting document format...")

        total_cores = os.cpu_count() or 8
        cpu_count = max(4, total_cores // 2)
        await self._send_log(log_queue, f"System detected {total_cores} CPU cores, using {cpu_count} cores")

        # Run conversion in thread
        def convert_worker():
            try:
                cv.convert(output_path, multi_processing=True, cpu_count=cpu_count)
            except TypeError:
                cv.convert(output_path)
            finally:
                cv.close()

        await asyncio.to_thread(convert_worker)
    
    async def convert(
        self,
        source_path: str,
        target_format: str,
        quality: str = 'high',
        task_id: str = None,
        log_queue: Optional[asyncio.Queue] = None,
        options: Optional[Dict[str, Any]] = None
    ) -> str:
        """Convert document to target format"""
        
        # Ensure os module is available in function scope
        import os
        
        # Determine source format from file extension
        source_format = Path(source_path).suffix.lower().lstrip('.')
        
        # Check if conversion is supported
        supported_targets = self.get_supported_targets(source_format)
        if target_format not in supported_targets:
            raise ConversionError(f"Conversion from {source_format} to {target_format} is not supported")
        
        # Generate unique conversion ID
        convert_id = str(uuid.uuid4())
        
        # Create temporary output file
        temp_dir = tempfile.gettempdir()
        output_filename = f"{Path(source_path).stem}.{target_format}"
        output_path = os.path.join(temp_dir, f"convert_{convert_id}_{output_filename}")
        
        # Store conversion task info
        conversion_tasks[convert_id] = {
            'task_id': task_id,
            'source_path': source_path,
            'target_format': target_format,
            'output_path': output_path,
            'intermediate_files': {},  # Will store md_path and html_path
            'status': 'processing',
            'start_time': datetime.now(),
            'error': None
        }
        
        # Store file info for cleanup
        conversion_files[convert_id] = {
            'output_path': output_path,
            'created_at': datetime.now(),
            'expires_at': datetime.now() + timedelta(minutes=30)  # 30 minutes retention
        }
        
        # Get options
        if not options:
            options = {}
        
        try:
            # Perform conversion based on format
            if source_format == 'pdf' and target_format == 'docx':
                # Check if MinerU should be used
                use_mineru = options.get('use_mineru', False)
                mineru_config = options.get('mineru_config', {})
                await self.convert_pdf_to_docx(source_path, output_path, quality, log_queue, use_mineru, mineru_config)
            else:
                raise ConversionError(f"Conversion from {source_format} to {target_format} not implemented")
            
            # Update task status
            conversion_tasks[convert_id]['status'] = 'completed'
            conversion_tasks[convert_id]['end_time'] = datetime.now()
            
            logger.info(f"Conversion completed: {convert_id}")
            return convert_id
            
        except Exception as e:
            # Update task status with error
            conversion_tasks[convert_id]['status'] = 'failed'
            conversion_tasks[convert_id]['error'] = str(e)
            conversion_tasks[convert_id]['end_time'] = datetime.now()
            
            # Clean up failed conversion file
            if os.path.exists(output_path):
                try:
                    os.remove(output_path)
                except:
                    pass
            
            logger.error(f"Conversion failed: {convert_id}, error: {e}")
            await self._send_log(log_queue, f"Conversion failed: {str(e)}")
            raise
    
    def get_conversion_status(self, convert_id: str) -> Dict[str, Any]:
        """Get conversion task status"""
        if convert_id not in conversion_tasks:
            raise ConversionError(f"Conversion task {convert_id} not found")
        
        task = conversion_tasks[convert_id]
        return {
            'convert_id': convert_id,
            'status': task['status'],
            'error': task.get('error'),
            'start_time': task['start_time'].isoformat(),
            'end_time': task.get('end_time').isoformat() if task.get('end_time') else None,
            'target_format': task.get('target_format', 'docx')
        }
    
    def get_conversion_file(self, convert_id: str) -> Optional[Dict[str, str]]:
        """Get converted file paths if conversion is completed

        Returns:
            Dict with keys: 'docx', 'md', 'html' containing file paths
            Returns None if conversion not found or not completed
        """
        if convert_id not in conversion_tasks:
            return None

        task = conversion_tasks[convert_id]
        if task['status'] != 'completed':
            return None

        output_path = task['output_path']
        if not os.path.exists(output_path):
            return None

        # Derive intermediate file paths from output_path naming pattern
        # output_path: /tmp/convert_{convert_id}_{source_stem}.docx
        # md_path: /tmp/convert_{convert_id}_{source_stem}_convert.md
        # html_path: /tmp/convert_{convert_id}_{source_stem}_convert.html
        base_name = os.path.basename(output_path).replace('.docx', '')
        temp_dir = os.path.dirname(output_path)
        md_path = os.path.join(temp_dir, f"{base_name}_convert.md")
        html_path = os.path.join(temp_dir, f"{base_name}_convert.html")

        result = {
            'docx': output_path,
            'md': md_path if os.path.exists(md_path) else None,
            'html': html_path if os.path.exists(html_path) else None
        }

        return result
    
    def cleanup_expired_files(self):
        """Clean up expired conversion files"""
        now = datetime.now()
        expired_convert_ids = []
        
        for convert_id, file_info in conversion_files.items():
            if now > file_info['expires_at']:
                expired_convert_ids.append(convert_id)
        
        for convert_id in expired_convert_ids:
            try:
                file_info = conversion_files[convert_id]
                output_path = file_info['output_path']
                
                # Remove file if it exists
                if os.path.exists(output_path):
                    os.remove(output_path)
                    logger.info(f"Cleaned up expired conversion file: {output_path}")
                
                # Remove from tracking
                del conversion_files[convert_id]
                if convert_id in conversion_tasks:
                    del conversion_tasks[convert_id]
                    
            except Exception as e:
                logger.error(f"Failed to cleanup conversion file {convert_id}: {e}")


# Global converter instance
converter = FormatConverter()


async def cleanup_task():
    """Background task to clean up expired files"""
    while True:
        try:
            converter.cleanup_expired_files()
        except Exception as e:
            logger.error(f"Cleanup task error: {e}")
        
        # Run cleanup every 5 minutes
        await asyncio.sleep(300)
