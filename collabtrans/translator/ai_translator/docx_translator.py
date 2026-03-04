# SPDX-FileCopyrightText: 2025 QinHan
# SPDX-License-Identifier: MPL-2.0
import asyncio
import os
import re
import tempfile
from dataclasses import dataclass
from io import BytesIO
from typing import Self, Literal, List, Dict, Any, Tuple

import docx
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.shared import Pt, RGBColor

from collabtrans.agents.segments_agent import SegmentsTranslateAgentConfig, SegmentsTranslateAgent
from collabtrans.ir.document import Document
from collabtrans.translator.ai_translator.base import AiTranslatorConfig, AiTranslator


@dataclass
class RunFormatInfo:
    """Run 格式信息，用于格式感知分片"""
    font_name: str | None = None
    font_size: int | None = None  # 单位：磅 (Pt)
    bold: bool | None = None
    italic: bool | None = None
    underline: bool | None = None
    color_rgb: str | None = None  # RGB 颜色值，格式："RRGGBB" 或 "RGBColor(r, g, b)"
    highlight_color: str | None = None  # 高亮颜色
    strikethrough: bool | None = None
    
    def __eq__(self, other) -> bool:
        """比较两个格式信息是否相同（用于分片判断）"""
        if not isinstance(other, RunFormatInfo):
            return False
        return (
            self.font_name == other.font_name and
            self.font_size == other.font_size and
            self.bold == other.bold and
            self.italic == other.italic and
            self.underline == other.underline and
            self.color_rgb == other.color_rgb and
            self.highlight_color == other.highlight_color and
            self.strikethrough == other.strikethrough
        )
    
    def __hash__(self) -> int:
        """用于字典键"""
        return hash((
            self.font_name, self.font_size, self.bold, self.italic,
            self.underline, self.color_rgb, self.highlight_color, self.strikethrough
        ))


def extract_run_format(run: Run) -> RunFormatInfo:
    """从 Run 对象提取格式信息"""
    font = run.font
    
    # 提取字体名称
    font_name = None
    if font.name:
        font_name = font.name
    
    # 提取字号
    font_size = None
    if font.size:
        font_size = font.size.pt
    
    # 提取颜色（RGB）
    color_rgb = None
    if font.color and font.color.rgb:
        # 转换为字符串表示
        rgb = font.color.rgb
        try:
            # 尝试多种方式提取 RGB 值
            r = g = b = None
            
            # 方法1: 尝试使用 r, g, b 属性
            if hasattr(rgb, 'r') and hasattr(rgb, 'g') and hasattr(rgb, 'b'):
                r, g, b = rgb.r, rgb.g, rgb.b
            # 方法2: 尝试使用 red, green, blue 属性
            elif hasattr(rgb, 'red') and hasattr(rgb, 'green') and hasattr(rgb, 'blue'):
                r, g, b = rgb.red, rgb.green, rgb.blue
            # 方法3: 尝试使用索引访问（如果是元组或列表）
            elif hasattr(rgb, '__getitem__') and len(rgb) >= 3:
                try:
                    r, g, b = rgb[0], rgb[1], rgb[2]
                except (TypeError, IndexError):
                    pass
            # 方法4: 尝试使用 hex() 方法转换为十六进制
            elif hasattr(rgb, 'hex'):
                hex_str = rgb.hex()
                if hex_str and len(hex_str) >= 6:
                    r = int(hex_str[0:2], 16)
                    g = int(hex_str[2:4], 16)
                    b = int(hex_str[4:6], 16)
            
            # 如果成功提取了 RGB 值，转换为字符串
            if r is not None and g is not None and b is not None:
                # 使用十六进制格式存储（更通用）
                color_rgb = f"{r:02X}{g:02X}{b:02X}"
            else:
                # 如果无法提取，使用字符串表示
                color_rgb = str(rgb)
        except Exception:
            # 如果所有方法都失败，使用字符串表示
            color_rgb = str(rgb)
    
    # 提取高亮颜色
    # 注意：某些第三方编辑器（如早期 WPS）可能写出 python-docx 不支持的值，比如 w:val="none"
    # 这会在访问 font.highlight_color 时触发 ValueError: WD_COLOR_INDEX has no XML mapping for 'none'
    # 完全避免使用 hasattr，直接使用 try/except 处理，因为 hasattr 在访问属性时可能无法捕获所有异常
    highlight_color = None
    try:
        # 直接访问 highlight_color 属性，不使用 hasattr
        hc = font.highlight_color
        if hc:
            highlight_color = str(hc)
    except ValueError as e:
        # 对于 "no XML mapping for 'none'" 这类已知兼容性问题，安全地忽略高亮信息，避免中断整个翻译任务
        msg = str(e)
        if "no XML mapping for 'none'" in msg or "has no XML mapping for 'none'" in msg:
            try:
                import logging
                logging.getLogger(__name__).warning(
                    "Ignored unsupported highlight_color 'none' in DOCX run: %s", msg
                )
            except Exception:
                # 日志失败不影响主流程
                pass
            highlight_color = None
        else:
            # 其它未知的 ValueError 仍然抛出，避免掩盖真正的问题
            raise
    except (AttributeError, TypeError):
        # 如果 font 对象没有 highlight_color 属性或访问时出现其他类型错误，安全忽略
        highlight_color = None
    
    # 提取删除线
    strikethrough = None
    if hasattr(font, 'strike'):
        strikethrough = font.strike
    
    return RunFormatInfo(
        font_name=font_name,
        font_size=font_size,
        bold=font.bold,
        italic=font.italic,
        underline=font.underline,
        color_rgb=color_rgb,
        highlight_color=highlight_color,
        strikethrough=strikethrough
    )


def apply_format_to_run(run: Run, format_info: RunFormatInfo, target_font_name: str | None = None, logger=None):
    """
    将格式信息应用到 Run 对象
    
    Args:
        run: Run 对象
        format_info: 格式信息
        target_font_name: 目标语言字体名称（可选，用于字体兼容性）
        logger: 日志记录器（可选）
    """
    if logger is None:
        import logging
        logger = logging.getLogger(__name__)
    
    font = run.font
    
    # 应用字体名称（优先使用目标语言字体，如果未指定则使用原始字体）
    if target_font_name:
        font.name = target_font_name
    elif format_info.font_name:
        font.name = format_info.font_name
    
    # 应用字号
    if format_info.font_size:
        font.size = Pt(format_info.font_size)
    
    # 应用粗体
    if format_info.bold is not None:
        font.bold = format_info.bold
    
    # 应用斜体
    if format_info.italic is not None:
        font.italic = format_info.italic
    
    # 应用下划线
    if format_info.underline is not None:
        font.underline = format_info.underline
    
    # 应用颜色
    if format_info.color_rgb:
        try:
            # 解析 RGB 颜色字符串
            r = g = b = None
            
            # 方法1: 尝试解析 "RGBColor(r, g, b)" 格式
            if format_info.color_rgb.startswith("RGBColor"):
                match = re.search(r'RGBColor\((\d+),\s*(\d+),\s*(\d+)\)', format_info.color_rgb)
                if match:
                    r, g, b = map(int, match.groups())
            # 方法2: 尝试解析十六进制格式 "RRGGBB"（6位）
            elif len(format_info.color_rgb) == 6:
                try:
                    r = int(format_info.color_rgb[0:2], 16)
                    g = int(format_info.color_rgb[2:4], 16)
                    b = int(format_info.color_rgb[4:6], 16)
                except ValueError:
                    pass
            # 方法3: 尝试解析十进制格式（如果包含逗号）
            elif ',' in format_info.color_rgb:
                parts = format_info.color_rgb.split(',')
                if len(parts) == 3:
                    try:
                        r, g, b = map(int, parts)
                    except ValueError:
                        pass
            
            # 如果成功解析了 RGB 值，应用颜色
            if r is not None and g is not None and b is not None:
                font.color.rgb = RGBColor(r, g, b)
        except Exception as e:
            logger.warning(f"Failed to apply color {format_info.color_rgb}: {e}")
    
    # 应用高亮颜色
    if format_info.highlight_color:
        try:
            # 注意：python-docx 对高亮颜色的支持可能有限
            # 这里需要根据实际情况调整
            if hasattr(font, 'highlight_color'):
                # 尝试解析高亮颜色
                pass
        except Exception as e:
            logger.warning(f"Failed to apply highlight color: {e}")
    
    # 应用删除线
    if format_info.strikethrough is not None:
        if hasattr(font, 'strike'):
            font.strike = format_info.strikethrough


def is_image_run(run: Run) -> bool:
    """Check if a run contains an image."""
    # w:drawing is the flag for embedded images, w:pict is the flag for VML images
    return '<w:drawing' in run.element.xml or '<w:pict' in run.element.xml


def has_page_break(run: Run) -> bool:
    """Check if a run contains a page break."""
    return '<w:br' in run._element.xml and 'w:type="page"' in run._element.xml


def preserve_page_breaks_in_run(run: Run, new_text: str) -> None:
    """Set run text while preserving page breaks."""
    # Check if run has page breaks
    breaks = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')
    page_breaks = [br for br in breaks if br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page']
    
    if page_breaks:
        # Clear existing text elements but keep breaks
        text_elements = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
        for text_elem in text_elements:
            text_elem.text = new_text
            # Only keep the first text element, remove others
            break
        else:
            # No text elements found, create one
            from docx.oxml import OxmlElement
            t_elem = OxmlElement('w:t')
            t_elem.text = new_text
            run._element.append(t_elem)
    else:
        # No page breaks, safe to set text directly
        run.text = new_text


def _paragraph_has_toc_field(paragraph: Paragraph) -> bool:
    """Check if a paragraph contains a TOC field."""
    try:
        p = paragraph._p  # lxml element
        
        # Check for TOC field codes
        fldChars = p.xpath('.//*[local-name()="fldChar"]')
        if not fldChars:
            # quick check for instruction text
            instrs = p.xpath('.//*[local-name()="instrText"]')
            for it in instrs:
                if 'TOC' in (it.text or ''):
                    return True
        else:
            instrs = p.xpath('.//*[local-name()="instrText"]')
            for it in instrs:
                if 'TOC' in (it.text or ''):
                    return True
                    
    except Exception:
        return False
    return False


def get_font_for_language(target_language: str) -> str:
    """
    Return appropriate font based on target language.
    
    Args:
        target_language: Target language name
        
    Returns:
        Recommended font name
    """
    # Language to font mapping
    language_font_map = {
        # Chinese
        "Chinese": "Microsoft YaHei",  # Microsoft YaHei
        "Simplified Chinese": "Microsoft YaHei",
        "Traditional Chinese": "Microsoft JhengHei",  # Microsoft JhengHei
        
        # English
        "English": "Calibri",
        
        # Japanese
        "Japanese": "Yu Gothic",  # Yu Gothic
        "Japanese": "Yu Gothic",
        "Japanese": "Yu Gothic",
        
        # Korean
        "Korean": "Malgun Gothic",  # Malgun Gothic
        "한국어": "Malgun Gothic",
        "Korean": "Malgun Gothic",
        
        # Russian
        "Russian": "Times New Roman",  # Common Russian font
        "Русский": "Times New Roman",
        "Russian": "Times New Roman",
        
        # Arabic
        "Arabic": "Arial Unicode MS",  # Supports Arabic characters
        "العَرَبِيَّة": "Arial Unicode MS",
        "Arabic": "Arial Unicode MS",
        
        # Other European languages
        "Spanish": "Calibri",
        "Español": "Calibri",
        "Spanish": "Calibri",
        
        "French": "Calibri",
        "Français": "Calibri",
        "French": "Calibri",
        
        "German": "Calibri",
        "Deutsch": "Calibri",
        "German": "Calibri",
        
        "Portuguese": "Calibri",
        "Português": "Calibri",
        "Portuguese": "Calibri",
        
        # Vietnamese
        "Vietnamese": "Arial Unicode MS",  # Supports Vietnamese characters
        "tiếng Việt": "Arial Unicode MS",
        "Vietnamese": "Arial Unicode MS",
        
        # Hebrew
        "Hebrew": "Arial Unicode MS",  # Supports Hebrew characters
        "Hebrew": "Arial Unicode MS",
        "עברית": "Arial Unicode MS",
        
        # Thai
        "Thai": "Arial Unicode MS",  # Supports Thai characters
        "Thai": "Arial Unicode MS",
        "ไทย": "Arial Unicode MS",
        
        # Hindi
        "Hindi": "Arial Unicode MS",  # Supports Hindi characters
        "Hindi": "Arial Unicode MS",
        "हिन्दी": "Arial Unicode MS",
    }
    
    # Find matching font
    font = language_font_map.get(target_language)
    
    # If no matching font found, return default font
    if not font:
        # Intelligent selection based on language characteristics
        if any(char in target_language for char in ['Chinese', 'Chinese', 'Simplified', 'Traditional']):
            font = "Microsoft YaHei"
        elif any(char in target_language for char in ['Japanese', 'Japanese', 'Japanese']):
            font = "Yu Gothic"
        elif any(char in target_language for char in ['Korean', 'Korean', '한국어']):
            font = "Malgun Gothic"
        elif any(char in target_language for char in ['Russian', 'Russian', 'Русский']):
            font = "Times New Roman"
        elif any(char in target_language for char in ['Arabic', 'Arabic', 'العَرَبِيَّة']):
            font = "Arial Unicode MS"
        elif any(char in target_language for char in ['Vietnamese', 'Vietnamese', 'tiếng Việt']):
            font = "Arial Unicode MS"
        elif any(char in target_language for char in ['Hebrew', 'Hebrew', 'עברית']):
            font = "Arial Unicode MS"
        elif any(char in target_language for char in ['Thai', 'Thai', 'ไทย']):
            font = "Arial Unicode MS"
        elif any(char in target_language for char in ['Hindi', 'Hindi', 'हिन्दी']):
            font = "Arial Unicode MS"
        else:
            # Default to Calibri (suitable for most Latin alphabet languages)
            font = "Calibri"
    
    return font


@dataclass
class DocxTranslatorConfig(AiTranslatorConfig):
    """
    Configuration class for DocxTranslator.
    """
    insert_mode: Literal["replace", "append", "prepend"] = "replace"
    separator: str = "\n"


class DocxTranslator(AiTranslator):
    """
    Translator for .docx files.
    This version is optimized to handle mixed text and image paragraphs without losing images.
    """

    def __init__(self, config: DocxTranslatorConfig):
        super().__init__(config=config)
        self.chunk_size = config.chunk_size
        self.translate_agent = None
        if not self.skip_translate:
            agent_config = SegmentsTranslateAgentConfig(
                custom_prompt=config.custom_prompt,
                to_lang=config.to_lang,
                base_url=config.base_url,
                api_key=config.api_key,
                api_type=getattr(config, 'api_type', 'openai'),
                model_id=config.model_id,
                temperature=config.temperature,
                thinking=config.thinking,
                concurrent=config.concurrent,
                timeout=config.timeout,
                logger=self.logger,
                glossary_dict=config.glossary_dict,
                retry=config.retry
            )
            self.translate_agent = SegmentsTranslateAgent(agent_config)
        self.insert_mode = config.insert_mode
        self.separator = config.separator

    def _pre_translate(self, document: Document) -> Tuple[DocumentObject, List[Dict[str, Any]], List[str]]:
        """
        [Refactored] Preprocess .docx file, extract text at Run level to avoid breaking images.
        :param document: Document object containing .docx file content.
        :return: A tuple containing:
                 - docx.Document object
                 - A list containing text block information (each element represents a group of consecutive text runs)
                 - A list containing all original texts to be translated
        """
        doc = docx.Document(BytesIO(document.content))
        elements_to_translate = []
        original_texts = []
        skipped_toc = 0

        def process_paragraph(para: Paragraph):
            nonlocal elements_to_translate, original_texts, skipped_toc
            
            # Skip paragraphs that contain TOC fields
            if _paragraph_has_toc_field(para):
                skipped_toc += 1
                try:
                    snippet = (para.text or "")[:120]
                except Exception:
                    snippet = ""
                try:
                    if getattr(self, "logger", None):
                        self.logger.info(f"[TOC] Skipping TOC paragraph: '{snippet}'")
                except Exception:
                    pass
                return
            
            # 格式感知分片：根据格式差异进行分片
            current_text_segment = ""
            current_runs = []
            current_format: RunFormatInfo | None = None

            for run in para.runs:
                if is_image_run(run):
                    # 遇到图片，结束当前分片
                    if current_text_segment.strip():
                        elements_to_translate.append({
                            "type": "text_runs",
                            "runs": current_runs,
                            "format": current_format  # 记录格式信息
                        })
                        original_texts.append(current_text_segment)
                    # 重置
                    current_text_segment = ""
                    current_runs = []
                    current_format = None
                else:
                    # 提取当前 run 的格式
                    run_format = extract_run_format(run)
                    
                    # 判断是否需要分片
                    # 条件：格式发生变化 或 当前分片为空
                    if current_format is None:
                        # 第一个 run，开始新分片
                        current_format = run_format
                        current_runs.append(run)
                        current_text_segment += run.text
                    elif current_format == run_format:
                        # 格式相同，继续累积
                        current_runs.append(run)
                        current_text_segment += run.text
                    else:
                        # 格式不同，结束当前分片，开始新分片
                        if current_text_segment.strip():
                            elements_to_translate.append({
                                "type": "text_runs",
                                "runs": current_runs,
                                "format": current_format
                            })
                            original_texts.append(current_text_segment)
                        
                        # 开始新分片
                        current_format = run_format
                        current_runs = [run]
                        current_text_segment = run.text

            # 处理最后一个分片
            if current_text_segment.strip():
                elements_to_translate.append({
                    "type": "text_runs",
                    "runs": current_runs,
                    "format": current_format
                })
                original_texts.append(current_text_segment)

        # Traverse all paragraphs
        for para in doc.paragraphs:
            process_paragraph(para)

        # Traverse all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        process_paragraph(para)

        # Log summary for TOC skipping
        try:
            if getattr(self, "logger", None):
                self.logger.info(f"[TOC] Skipped paragraphs counted as TOC: {skipped_toc}")
        except Exception:
            pass

        return doc, elements_to_translate, original_texts

    def _after_translate(self, doc: DocumentObject, elements_to_translate: List[Dict[str, Any]],
                         translated_texts: List[str], original_texts: List[str]) -> bytes:
        """
        [Refactored] Write translated text back to corresponding text runs, preserving images and styles.
        格式感知版本：保留原始格式属性（字号、颜色、粗体等）
        """
        translation_map = dict(zip(original_texts, translated_texts))
        
        # 获取目标语言字体（用于字体兼容性）
        target_font_name = get_font_for_language(self.config.to_lang)

        for i, element_info in enumerate(elements_to_translate):
            runs = element_info["runs"]
            format_info: RunFormatInfo | None = element_info.get("format")
            original_text = original_texts[i]
            translated_text = translated_texts[i]

            # Determine final text based on insert mode
            if self.insert_mode == "replace":
                final_text = translated_text
            elif self.insert_mode == "append":
                final_text = original_text + self.separator + translated_text
            elif self.insert_mode == "prepend":
                final_text = translated_text + self.separator + original_text
            else:
                self.logger.error("Invalid DocxTranslatorConfig parameter")
                final_text = translated_text

            if not runs:
                continue

            # --- 格式感知导出逻辑 ---
            # 策略：将翻译文本分配到各个 run，保持格式
            if len(runs) == 1:
                # 单个 run：直接替换并应用格式
                first_run = runs[0]
                preserve_page_breaks_in_run(first_run, final_text)
                if format_info:
                    apply_format_to_run(first_run, format_info, target_font_name, self.logger)
                else:
                    # 如果没有格式信息，使用原有逻辑（字体兼容性）
                    if first_run.font:
                        first_run.font.name = target_font_name
            else:
                # 多个 run：将文本放入第一个 run，其他 run 清空
                first_run = runs[0]
                preserve_page_breaks_in_run(first_run, final_text)
                if format_info:
                    apply_format_to_run(first_run, format_info, target_font_name, self.logger)
                else:
                    # 如果没有格式信息，使用原有逻辑（字体兼容性）
                    if first_run.font:
                        first_run.font.name = target_font_name
                        # 尝试备选字体
                        if not first_run.font.name:
                            if any(char in self.config.to_lang for char in ['Chinese', 'Chinese', 'Simplified', 'Traditional']):
                                fallback_fonts = ['SimSun', 'SimHei', 'Arial Unicode MS', 'Times New Roman']
                            elif any(char in self.config.to_lang for char in ['Japanese', 'Japanese', 'Japanese']):
                                fallback_fonts = ['MS Gothic', 'Arial Unicode MS', 'Times New Roman']
                            elif any(char in self.config.to_lang for char in ['Korean', 'Korean', '한국어']):
                                fallback_fonts = ['Gulim', 'Arial Unicode MS', 'Times New Roman']
                            elif any(char in self.config.to_lang for char in ['Russian', 'Russian', 'Русский']):
                                fallback_fonts = ['Times New Roman', 'Arial', 'Calibri']
                            elif any(char in self.config.to_lang for char in ['Arabic', 'Arabic', 'العَرَبِيَّة']):
                                fallback_fonts = ['Arial Unicode MS', 'Times New Roman', 'Arial']
                            else:
                                fallback_fonts = ['Calibri', 'Times New Roman', 'Arial']
                            
                            for fallback_font in fallback_fonts:
                                first_run.font.name = fallback_font
                                if first_run.font.name:
                                    break
                
                # 清空其他 run（保留结构）
                for run in runs[1:]:
                    run.text = ""
                    # 可选：也应用格式（保持一致性）
                    if format_info:
                        apply_format_to_run(run, format_info, target_font_name, self.logger)
            # --- End of format-aware logic ---

        # Save to temp file to reduce peak memory (avoid BytesIO buffer + getvalue() copy for large docx)
        fd, path = tempfile.mkstemp(suffix=".docx")
        try:
            os.close(fd)
            doc.save(path)
            with open(path, "rb") as f:
                return f.read()
        finally:
            try:
                os.unlink(path)
            except OSError:
                pass

    def translate(self, document: Document) -> Self:
        """
        Synchronously translate .docx file.
        """
        doc, elements_to_translate, original_texts = self._pre_translate(document)
        document.content = b""  # Free original file bytes early to reduce peak memory
        if not original_texts:
            # Use i18n logger for translation messages
            from collabtrans.logger.logger import i18n_logger
            i18n_logger.info("backend.translation.task.no_text_found")
            output_stream = BytesIO()
            doc.save(output_stream)
            document.content = output_stream.getvalue()
            return self

        if self.glossary_agent:
            self.glossary_dict_gen = self.glossary_agent.send_segments(original_texts, self.chunk_size)
            if self.translate_agent:
                self.translate_agent.update_glossary_dict(self.glossary_dict_gen)

        # Call translation agent
        if self.translate_agent:
            translated_texts = self.translate_agent.send_segments(original_texts, self.chunk_size)
        else:
            translated_texts = original_texts

        # Write translation results back to document
        document.content = self._after_translate(doc, elements_to_translate, translated_texts, original_texts)
        return self

    async def translate_async(self, document: Document) -> Self:
        """
        Asynchronously translate .docx file.
        """
        doc, elements_to_translate, original_texts = await asyncio.to_thread(self._pre_translate, document)
        document.content = b""  # Free original file bytes early to reduce peak memory
        if not original_texts:
            # Use i18n logger for translation messages
            from collabtrans.logger.logger import i18n_logger
            i18n_logger.info("backend.translation.task.no_text_found")
            # Correctly save and return in async environment
            output_stream = BytesIO()
            doc.save(output_stream)
            document.content = output_stream.getvalue()
            return self

        if self.glossary_agent:
            self.glossary_dict_gen = await self.glossary_agent.send_segments_async(original_texts, self.chunk_size)
            if self.translate_agent:
                self.translate_agent.update_glossary_dict(self.glossary_dict_gen)

        # Asynchronously call translation agent
        if self.translate_agent:
            translated_texts = await self.translate_agent.send_segments_async(original_texts, self.chunk_size)
        else:
            translated_texts = original_texts
        # Write translation results back to document
        document.content = await asyncio.to_thread(self._after_translate, doc, elements_to_translate, translated_texts,
                                                   original_texts)
        return self
