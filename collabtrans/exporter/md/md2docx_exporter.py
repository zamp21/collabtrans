# SPDX-FileCopyrightText: 2025 QinHan
# SPDX-License-Identifier: MPL-2.0

import base64
import hashlib
import logging
import mimetypes
import os
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Optional, Tuple

from collabtrans.exporter.md.base import MDExporter, MDExporterConfig
from collabtrans.ir.document import Document
from collabtrans.ir.markdown_document import MarkdownDocument

logger = logging.getLogger(__name__)


@dataclass(kw_only=True)
class MD2DocxExporterConfig(MDExporterConfig):
    """Configuration for MD to DOCX exporter using Pandoc."""
    # Additional Pandoc options can be added here
    reference_doc: Optional[str] = None  # Path to reference DOCX for styling
    toc: bool = False  # Whether to include table of contents
    toc_depth: int = 3  # Depth of table of contents
    extract_images: bool = True  # Whether to extract base64 images to files


class MD2DocxExporter(MDExporter):
    """
    Export Markdown to DOCX using Pandoc.

    Uses a two-step conversion process:
    1. Markdown → HTML (preserves HTML tables from MinerU using raw_html extension)
    2. HTML → DOCX (converts HTML tables to proper DOCX tables)

    Handles:
    - Base64 embedded images (extracts to temporary files)
    - HTML tables (inline tables from MinerU OCR)
    - Markdown pipe tables
    - LaTeX formulas (preserved through conversion)

    Requires pandoc to be installed on the system.
    """

    def __init__(self, config: Optional[MD2DocxExporterConfig] = None):
        super().__init__(config=config)
        self._check_pandoc_available()

    def _check_pandoc_available(self):
        """Check if pandoc is available on the system."""
        try:
            import subprocess
            result = subprocess.run(['pandoc', '--version'], capture_output=True, timeout=5)
            if result.returncode != 0:
                raise RuntimeError("Pandoc is not available. Please install pandoc: https://pandoc.org/installing.html")
        except FileNotFoundError:
            raise RuntimeError("Pandoc is not installed. Please install pandoc: https://pandoc.org/installing.html")
        except Exception as e:
            raise RuntimeError(f"Failed to check pandoc availability: {e}")

    def _extract_base64_images(self, markdown_content: str, image_dir: str, md_dir: str) -> Tuple[str, list]:
        """
        Extract base64 embedded images to files and update markdown references.

        Args:
            markdown_content: The markdown content with base64 images
            image_dir: Directory to save extracted images
            md_dir: Directory where the markdown file will be saved (for relative paths)

        Returns:
            Tuple of (updated markdown content, list of created image paths)
        """
        os.makedirs(image_dir, exist_ok=True)
        created_files = []

        # Pattern to match base64 images: ![alt](data:mime;base64,data)
        pattern = r'!\[(.*?)\]\(data:([^;]+);base64,([^)]+)\)'

        def replace_base64_image(match: re.Match) -> str:
            alt_text = match.group(1)
            mime_type = match.group(2)
            b64_data = match.group(3)

            # Determine file extension
            ext = mimetypes.guess_extension(mime_type)
            if not ext:
                ext = '.bin'

            # Generate unique filename based on content hash
            content_hash = hashlib.md5(b64_data.encode()).hexdigest()[:12]
            image_filename = f"img_{content_hash}{ext}"
            image_path = os.path.join(image_dir, image_filename)

            # Decode and save image
            try:
                image_bytes = base64.b64decode(b64_data)
                with open(image_path, 'wb') as f:
                    f.write(image_bytes)
                created_files.append(image_path)
                logger.debug(f"Extracted base64 image to: {image_path}")

                # Calculate relative path from markdown file to image
                rel_image_path = os.path.relpath(image_path, md_dir)
                # Return updated markdown with relative file reference
                return f"![{alt_text}]({rel_image_path})"
            except Exception as e:
                logger.warning(f"Failed to extract base64 image: {e}")
                return match.group(0)  # Return original if extraction fails

        updated_content = re.sub(pattern, replace_base64_image, markdown_content)
        return updated_content, created_files

    def _process_html_tables(self, markdown_content: str) -> str:
        """
        Process HTML tables to ensure they render correctly in DOCX.

        Add border attribute to HTML tables so they have visible borders in DOCX.
        MinerU generates tables without border attributes, which results in
        invisible borders in the exported DOCX.
        """
        # Add border="1" to all <table> tags that don't already have border attribute
        # This ensures tables have visible borders in DOCX
        def add_border_to_table(match):
            table_tag = match.group(0)
            # Check if border attribute already exists
            if 'border' not in table_tag:
                return table_tag.replace('<table', '<table border="1"', 1)
            return table_tag

        # Match <table> tags (including those with attributes)
        updated_content = re.sub(r'<table\b[^>]*>', add_border_to_table, markdown_content)
        return updated_content

    def _add_table_border_style_to_html(self, html_content: str) -> str:
        """
        Add border attribute to HTML tables after MD→HTML conversion.

        Pandoc converts HTML tables with border="1" attribute to DOCX tables with borders.
        CSS styles don't work for Pandoc's HTML→DOCX conversion, so we use the border attribute.
        """
        # Add border="1" to all <table> tags that don't already have it
        def add_border(match):
            table_tag = match.group(0)
            if 'border=' not in table_tag.lower():
                # Add border="1" attribute
                return '<table border="1"' + table_tag[len('<table'):]
            return table_tag

        html_content = re.sub(r'<table\b[^>]*>', add_border, html_content)
        return html_content

    def export(self, document: MarkdownDocument) -> Document:
        """
        Export MarkdownDocument to DOCX format using Pandoc.

        Args:
            document: The MarkdownDocument to export

        Returns:
            A Document containing the DOCX file content
        """
        import subprocess

        config = self.config if isinstance(self.config, MD2DocxExporterConfig) else MD2DocxExporterConfig()

        # Create a temporary directory for all files
        # Use system temp directory and keep files for debugging
        temp_dir = tempfile.gettempdir()
        debug_prefix = f"translation_{document.stem}"
        md_path = os.path.join(temp_dir, f"{debug_prefix}_input.md")
        docx_path = os.path.join(temp_dir, f"{debug_prefix}_output.docx")
        html_path = os.path.join(temp_dir, f"{debug_prefix}_intermediate.html")
        image_dir = os.path.join(temp_dir, f"{debug_prefix}_images")

        # Get markdown content
        markdown_content = document.content.decode('utf-8')

        # Save original markdown for debugging
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        logger.info(f"[Translation DOCX] Saved input markdown to: {md_path}")

        # Count HTML tables in input
        html_tables_in_input = len(re.findall(r'<table[^>]*>.*?</table>', markdown_content, re.IGNORECASE | re.DOTALL))
        logger.info(f"[Translation DOCX] HTML tables in input markdown: {html_tables_in_input}")

        # Process content
        if config.extract_images:
            # Extract base64 images to files (use relative paths from md_dir)
            markdown_content, _ = self._extract_base64_images(markdown_content, image_dir, temp_dir)

        # Process HTML tables
        markdown_content = self._process_html_tables(markdown_content)

        # Save processed markdown for debugging
        processed_md_path = os.path.join(temp_dir, f"{debug_prefix}_processed.md")
        with open(processed_md_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        logger.info(f"[Translation DOCX] Saved processed markdown to: {processed_md_path}")

        try:
            # Two-step conversion to properly handle HTML tables:
            # 1. Markdown → HTML (preserves HTML tables from MinerU)
            # 2. HTML → DOCX (converts HTML tables to proper DOCX tables)

            # Step 1: Convert markdown to HTML
            # Using raw_html extension to preserve inline HTML tables
            md_to_html_cmd = [
                'pandoc',
                processed_md_path,
                '-o', html_path,
                '--from=markdown+pipe_tables+grid_tables+multiline_tables+raw_html+tex_math_dollars',
                '--to=html',
                '--resource-path=' + temp_dir,
            ]

            if config.reference_doc and os.path.exists(config.reference_doc):
                md_to_html_cmd.extend(['--reference-doc', config.reference_doc])

            logger.info(f"[Translation DOCX] Running pandoc MD→HTML: {' '.join(md_to_html_cmd)}")

            md_to_html_result = subprocess.run(
                md_to_html_cmd,
                capture_output=True,
                text=True,
                timeout=120
            )

            if md_to_html_result.returncode != 0:
                logger.error(f"Pandoc MD→HTML failed: {md_to_html_result.stderr}")
                raise RuntimeError(f"Pandoc MD→HTML conversion failed: {md_to_html_result.stderr}")

            # Read the generated HTML and add table border styles
            with open(html_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            # Count HTML tables in intermediate HTML
            html_tables_in_html = len(re.findall(r'<table[^>]*>.*?</table>', html_content, re.IGNORECASE | re.DOTALL))
            logger.info(f"[Translation DOCX] HTML tables in intermediate HTML: {html_tables_in_html}")

            # Add CSS border styles for tables
            html_content = self._add_table_border_style_to_html(html_content)

            # Write the modified HTML back
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)

            logger.info(f"[Translation DOCX] Saved intermediate HTML to: {html_path}")
            logger.info("[Translation DOCX] Added table border styles to HTML")

            # Step 2: Convert HTML to DOCX
            # This properly converts HTML tables to DOCX tables
            html_to_docx_cmd = [
                'pandoc',
                html_path,
                '-o', docx_path,
                '--from=html',
                '--to=docx',
                '--wrap=none',
                '--resource-path=' + temp_dir,
            ]

            if config.reference_doc and os.path.exists(config.reference_doc):
                html_to_docx_cmd.extend(['--reference-doc', config.reference_doc])

            if config.toc:
                html_to_docx_cmd.append('--toc')
                html_to_docx_cmd.extend(['--toc-depth', str(config.toc_depth)])

            logger.info(f"[Translation DOCX] Running pandoc HTML→DOCX: {' '.join(html_to_docx_cmd)}")

            html_to_docx_result = subprocess.run(
                html_to_docx_cmd,
                capture_output=True,
                text=True,
                timeout=180  # 3 minutes for HTML to DOCX
            )

            if html_to_docx_result.returncode != 0:
                logger.error(f"Pandoc HTML→DOCX failed: {html_to_docx_result.stderr}")
                raise RuntimeError(f"Pandoc HTML→DOCX conversion failed: {html_to_docx_result.stderr}")

            # Read the generated DOCX file
            with open(docx_path, 'rb') as f:
                docx_content = f.read()

            # Add table borders using python-docx (Pandoc doesn't add borders from HTML)
            try:
                from docx import Document as DocxDocument
                from docx.oxml import parse_xml
                from docx.oxml.ns import qn

                # Load the DOCX
                doc = DocxDocument(docx_path)

                logger.info(f"[Translation DOCX] Found {len(doc.tables)} tables in DOCX before border processing")

                # Add borders to all tables
                for table in doc.tables:
                    tbl = table._tbl
                    tblPr = tbl.tblPr

                    # Find or create tblPr
                    if tblPr is None:
                        tblPr = parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                        tbl.insert(0, tblPr)

                    # Remove tblStyle element to prevent style borders from overriding our borders
                    # Pandoc's "Table" style has conditional formatting with nil borders that override tblBorders
                    existing_style = tblPr.find(qn('w:tblStyle'))
                    if existing_style is not None:
                        tblPr.remove(existing_style)
                        logger.debug("Removed table style reference to enable explicit borders")

                    # Create tblBorders element with proper namespace
                    # sz=12 means 3/4 point (12/8 = 1.5 pt), visible borders
                    tblBorders_xml = parse_xml(
                        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        r'<w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                        r'<w:left w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                        r'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                        r'<w:right w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                        r'<w:insideH w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                        r'<w:insideV w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                        r'</w:tblBorders>'
                    )

                    # Remove existing tblBorders if present (find by tag name)
                    existing_borders = tblPr.find(qn('w:tblBorders'))
                    if existing_borders is not None:
                        tblPr.remove(existing_borders)

                    # Add new borders to tblPr
                    tblPr.append(tblBorders_xml)

                # Save the modified DOCX
                doc.save(docx_path)

                # Read the modified content
                with open(docx_path, 'rb') as f:
                    docx_content = f.read()

                logger.info(f"[Translation DOCX] Added table borders to DOCX (found {len(doc.tables)} tables)")
                logger.info(f"[Translation DOCX] Saved final DOCX to: {docx_path}")

            except ImportError:
                logger.warning("python-docx not available, tables may not have borders. Install: pip install python-docx")
            except Exception as e:
                logger.warning(f"Failed to add table borders: {e}")

            logger.info(f"[Translation DOCX] Successfully converted markdown to DOCX ({len(docx_content)} bytes)")

            return Document.from_bytes(
                suffix=".docx",
                content=docx_content,
                stem=document.stem
            )

        except subprocess.TimeoutExpired:
            raise RuntimeError("Pandoc conversion timed out")