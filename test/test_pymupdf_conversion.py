#!/usr/bin/env python3
"""
PyMuPDF PDF to DOCX conversion test script
直接使用PyMuPDF进行PDF转DOCX转换测试
"""

import fitz  # PyMuPDF
import time
import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from PIL import Image

def test_pdf_text_extraction(pdf_path):
    """测试PDF文本提取"""
    print(f"🔍 测试PDF文本提取: {pdf_path}")
    
    try:
        doc = fitz.open(pdf_path)
        total_text = ""
        page_count = len(doc)
        
        print(f"📄 PDF页数: {page_count}")
        
        for page_num in range(min(3, page_count)):  # 只检查前3页
            page = doc[page_num]
            text = page.get_text()
            total_text += text
            print(f"   第{page_num + 1}页文本长度: {len(text)} 字符")
        
        doc.close()
        
        print(f"📊 总文本长度: {len(total_text)} 字符")
        print(f"📝 前100字符预览: {total_text[:100]}...")
        
        return len(total_text.strip()) > 50, total_text
        
    except Exception as e:
        print(f"❌ 文本提取失败: {e}")
        return False, ""

def test_pdf_image_extraction(pdf_path):
    """测试PDF图片提取"""
    print(f"🖼️ 测试PDF图片提取: {pdf_path}")
    
    try:
        doc = fitz.open(pdf_path)
        total_images = 0
        
        for page_num in range(min(3, len(doc))):  # 只检查前3页
            page = doc[page_num]
            image_list = page.get_images()
            total_images += len(image_list)
            print(f"   第{page_num + 1}页图片数量: {len(image_list)}")
        
        doc.close()
        print(f"📊 总图片数量: {total_images}")
        return total_images
        
    except Exception as e:
        print(f"❌ 图片提取失败: {e}")
        return 0

def convert_pdf_to_docx_pymupdf(pdf_path, output_path):
    """使用PyMuPDF直接转换PDF到DOCX"""
    print(f"🚀 开始PyMuPDF转换: {pdf_path} -> {output_path}")
    
    start_time = time.time()
    
    try:
        # 打开PDF文档
        doc = fitz.open(pdf_path)
        page_count = len(doc)
        print(f"📄 处理 {page_count} 页")
        
        # 创建Word文档
        word_doc = Document()
        
        for page_num in range(page_count):
            page = doc[page_num]
            print(f"   处理第 {page_num + 1} 页...")
            
            # 提取文本
            text = page.get_text()
            if text.strip():
                # 添加文本到Word文档
                paragraph = word_doc.add_paragraph(text)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 提取图片
            image_list = page.get_images()
            for img_index, img in enumerate(image_list):
                try:
                    # 获取图片数据
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    
                    if pix.n - pix.alpha < 4:  # 确保不是CMYK
                        # 转换为PIL Image
                        img_data = pix.tobytes("png")
                        pil_image = Image.open(io.BytesIO(img_data))
                        
                        # 保存临时图片
                        temp_img_path = f"/tmp/temp_img_{page_num}_{img_index}.png"
                        pil_image.save(temp_img_path)
                        
                        # 添加到Word文档
                        word_doc.add_picture(temp_img_path, width=Inches(4))
                        
                        # 清理临时文件
                        os.remove(temp_img_path)
                        
                except Exception as e:
                    print(f"   图片 {img_index} 处理失败: {e}")
                    continue
            
            # 添加分页符（除了最后一页）
            if page_num < page_count - 1:
                word_doc.add_page_break()
        
        # 保存Word文档
        word_doc.save(output_path)
        doc.close()
        
        end_time = time.time()
        duration = end_time - start_time
        
        print(f"✅ 转换完成! 耗时: {duration:.2f} 秒")
        print(f"📁 输出文件: {output_path}")
        
        # 检查输出文件大小
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            print(f"📊 输出文件大小: {file_size / 1024:.2f} KB")
        
        return True, duration
        
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        return False, 0

def test_pdf2docx_comparison(pdf_path):
    """对比pdf2docx和PyMuPDF的转换效果"""
    print(f"🔄 对比测试: {pdf_path}")
    
    # 测试pdf2docx
    print("\n" + "="*50)
    print("📦 测试 pdf2docx")
    print("="*50)
    
    try:
        from pdf2docx import Converter
        
        output_pdf2docx = "/tmp/test_pdf2docx_output.docx"
        start_time = time.time()
        
        cv = Converter(pdf_path)
        cv.convert(output_pdf2docx)
        cv.close()
        
        pdf2docx_time = time.time() - start_time
        print(f"✅ pdf2docx 转换完成，耗时: {pdf2docx_time:.2f} 秒")
        
    except Exception as e:
        print(f"❌ pdf2docx 转换失败: {e}")
        pdf2docx_time = 0
    
    # 测试PyMuPDF
    print("\n" + "="*50)
    print("🔧 测试 PyMuPDF")
    print("="*50)
    
    output_pymupdf = "/tmp/test_pymupdf_output.docx"
    success, pymupdf_time = convert_pdf_to_docx_pymupdf(pdf_path, output_pymupdf)
    
    # 对比结果
    print("\n" + "="*50)
    print("📊 对比结果")
    print("="*50)
    print(f"pdf2docx 耗时: {pdf2docx_time:.2f} 秒")
    print(f"PyMuPDF 耗时: {pymupdf_time:.2f} 秒")
    
    if pdf2docx_time > 0 and pymupdf_time > 0:
        if pymupdf_time < pdf2docx_time:
            improvement = ((pdf2docx_time - pymupdf_time) / pdf2docx_time) * 100
            print(f"🚀 PyMuPDF 比 pdf2docx 快 {improvement:.1f}%")
        else:
            slowdown = ((pymupdf_time - pdf2docx_time) / pdf2docx_time) * 100
            print(f"⚠️ PyMuPDF 比 pdf2docx 慢 {slowdown:.1f}%")

def main():
    """主函数"""
    print("🧪 PyMuPDF PDF转DOCX测试脚本")
    print("="*50)
    
    # 检查命令行参数
    if len(sys.argv) < 2:
        print("❌ 用法: python test_pymupdf_conversion.py <PDF文件路径>")
        print("📝 示例: python test_pymupdf_conversion.py /path/to/test.pdf")
        sys.exit(1)
    
    pdf_path = sys.argv[1]
    
    if not os.path.exists(pdf_path):
        print(f"❌ PDF文件不存在: {pdf_path}")
        sys.exit(1)
    
    print(f"📁 测试文件: {pdf_path}")
    print(f"📊 文件大小: {os.path.getsize(pdf_path) / 1024:.2f} KB")
    
    # 1. 测试PDF文本提取
    print("\n" + "="*50)
    print("🔍 步骤1: PDF文本提取测试")
    print("="*50)
    has_text, text_content = test_pdf_text_extraction(pdf_path)
    
    # 2. 测试PDF图片提取
    print("\n" + "="*50)
    print("🖼️ 步骤2: PDF图片提取测试")
    print("="*50)
    image_count = test_pdf_image_extraction(pdf_path)
    
    # 3. 判断PDF类型
    print("\n" + "="*50)
    print("📋 步骤3: PDF类型判断")
    print("="*50)
    if has_text:
        print("✅ 文本版PDF - 包含可提取文本")
    else:
        print("📷 扫描版PDF - 主要包含图片")
    
    print(f"🖼️ 图片数量: {image_count}")
    
    # 4. 执行转换测试
    print("\n" + "="*50)
    print("🚀 步骤4: 转换测试")
    print("="*50)
    test_pdf2docx_comparison(pdf_path)
    
    print("\n" + "="*50)
    print("✅ 测试完成!")
    print("="*50)

if __name__ == "__main__":
    main()
