#!/usr/bin/env python3
"""
pdfminer PDF to DOCX conversion test script
使用pdfminer进行PDF转DOCX转换测试
"""

import time
import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
from PIL import Image

# pdfminer imports
from pdfminer.high_level import extract_text, extract_pages
from pdfminer.layout import LTTextContainer, LTChar, LTTextLine, LTTextBox, LTFigure, LTImage
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import resolve1
from pdfminer.psparser import PSKeyword, PSLiteral
from pdfminer.utils import open_filename

def test_pdfminer_text_extraction(pdf_path):
    """测试pdfminer文本提取"""
    print(f"🔍 测试pdfminer文本提取: {pdf_path}")
    
    try:
        # 使用高级API提取文本
        text = extract_text(pdf_path)
        print(f"📊 总文本长度: {len(text)} 字符")
        print(f"📝 前100字符预览: {text[:100]}...")
        
        # 使用低级API获取更详细信息
        with open_filename(pdf_path, "rb") as fp:
            parser = PDFParser(fp)
            doc = PDFDocument(parser)
            
            # 获取页面信息
            pages = list(PDFPage.create_pages(doc))
            print(f"📄 PDF页数: {len(pages)}")
            
            # 分析每页的文本
            rsrcmgr = PDFResourceManager()
            laparams = None
            device = PDFPageAggregator(rsrcmgr, laparams=laparams)
            interpreter = PDFPageInterpreter(rsrcmgr, device)
            
            page_text_lengths = []
            for page_num, page in enumerate(pages[:3]):  # 只检查前3页
                interpreter.process_page(page)
                layout = device.get_result()
                
                page_text = ""
                for element in layout:
                    if isinstance(element, LTTextContainer):
                        page_text += element.get_text()
                
                page_text_lengths.append(len(page_text))
                print(f"   第{page_num + 1}页文本长度: {len(page_text)} 字符")
        
        return len(text.strip()) > 50, text, page_text_lengths
        
    except Exception as e:
        print(f"❌ 文本提取失败: {e}")
        return False, "", []

def test_pdfminer_layout_analysis(pdf_path):
    """测试pdfminer布局分析"""
    print(f"📐 测试pdfminer布局分析: {pdf_path}")
    
    try:
        with open_filename(pdf_path, "rb") as fp:
            parser = PDFParser(fp)
            doc = PDFDocument(parser)
            pages = list(PDFPage.create_pages(doc))
            
            rsrcmgr = PDFResourceManager()
            laparams = None
            device = PDFPageAggregator(rsrcmgr, laparams=laparams)
            interpreter = PDFPageInterpreter(rsrcmgr, device)
            
            total_text_boxes = 0
            total_figures = 0
            total_images = 0
            
            for page_num, page in enumerate(pages[:3]):  # 只检查前3页
                interpreter.process_page(page)
                layout = device.get_result()
                
                page_text_boxes = 0
                page_figures = 0
                page_images = 0
                
                for element in layout:
                    if isinstance(element, LTTextBox):
                        page_text_boxes += 1
                    elif isinstance(element, LTFigure):
                        page_figures += 1
                        # 检查图片
                        for sub_element in element:
                            if isinstance(sub_element, LTImage):
                                page_images += 1
                
                total_text_boxes += page_text_boxes
                total_figures += page_figures
                total_images += page_images
                
                print(f"   第{page_num + 1}页: {page_text_boxes}文本框, {page_figures}图形, {page_images}图片")
            
            print(f"📊 总计: {total_text_boxes}文本框, {total_figures}图形, {total_images}图片")
            return total_text_boxes, total_figures, total_images
            
    except Exception as e:
        print(f"❌ 布局分析失败: {e}")
        return 0, 0, 0

def extract_text_with_formatting(pdf_path):
    """提取带格式的文本"""
    try:
        with open_filename(pdf_path, "rb") as fp:
            parser = PDFParser(fp)
            doc = PDFDocument(parser)
            pages = list(PDFPage.create_pages(doc))
            
            rsrcmgr = PDFResourceManager()
            laparams = None
            device = PDFPageAggregator(rsrcmgr, laparams=laparams)
            interpreter = PDFPageInterpreter(rsrcmgr, device)
            
            formatted_text = []
            
            for page in pages:
                interpreter.process_page(page)
                layout = device.get_result()
                
                page_text = []
                for element in layout:
                    if isinstance(element, LTTextContainer):
                        # 提取文本行
                        for line in element:
                            if isinstance(line, LTTextLine):
                                line_text = ""
                                for char in line:
                                    if isinstance(char, LTChar):
                                        line_text += char.get_text()
                                if line_text.strip():
                                    page_text.append(line_text.strip())
                
                formatted_text.extend(page_text)
            
            return formatted_text
            
    except Exception as e:
        print(f"❌ 格式化文本提取失败: {e}")
        return []

def convert_pdf_to_docx_pdfminer(pdf_path, output_path, preserve_formatting=True):
    """使用pdfminer转换PDF到DOCX"""
    print(f"🚀 开始pdfminer转换: {pdf_path} -> {output_path}")
    
    start_time = time.time()
    
    try:
        # 创建Word文档
        word_doc = Document()
        
        if preserve_formatting:
            # 使用格式化提取
            formatted_text = extract_text_with_formatting(pdf_path)
            
            for line in formatted_text:
                if line.strip():
                    paragraph = word_doc.add_paragraph(line)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # 使用简单文本提取
            text = extract_text(pdf_path)
            if text.strip():
                # 按行分割文本
                lines = text.split('\n')
                for line in lines:
                    line = line.strip()
                    if line:
                        paragraph = word_doc.add_paragraph(line)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 保存Word文档
        word_doc.save(output_path)
        
        end_time = time.time()
        duration = end_time - start_time
        
        print(f"✅ 转换完成! 耗时: {duration:.2f} 秒")
        print(f"📁 输出文件: {output_path}")
        
        # 检查输出文件大小
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            print(f"📊 输出文件大小: {file_size / 1024:.2f} KB")
        
        return True, duration, file_size if os.path.exists(output_path) else 0
        
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        return False, 0, 0

def test_pdfminer_vs_others(pdf_path):
    """对比pdfminer和其他转换方法"""
    print(f"🔄 对比测试: {pdf_path}")
    
    # 测试pdfminer (格式化)
    print("\n" + "="*50)
    print("📦 测试 pdfminer (格式化)")
    print("="*50)
    
    output_pdfminer_formatted = "/tmp/test_pdfminer_formatted_output.docx"
    success1, pdfminer_time1, size1 = convert_pdf_to_docx_pdfminer(pdf_path, output_pdfminer_formatted, preserve_formatting=True)
    
    # 测试pdfminer (简单)
    print("\n" + "="*50)
    print("📦 测试 pdfminer (简单)")
    print("="*50)
    
    output_pdfminer_simple = "/tmp/test_pdfminer_simple_output.docx"
    success2, pdfminer_time2, size2 = convert_pdf_to_docx_pdfminer(pdf_path, output_pdfminer_simple, preserve_formatting=False)
    
    # 测试pdfplumber (如果可用)
    print("\n" + "="*50)
    print("📦 测试 pdfplumber")
    print("="*50)
    
    try:
        import pdfplumber
        
        output_pdfplumber = "/tmp/test_pdfplumber_output.docx"
        start_time = time.time()
        
        with pdfplumber.open(pdf_path) as pdf:
            word_doc = Document()
            
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    lines = text.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:
                            paragraph = word_doc.add_paragraph(line)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            word_doc.save(output_pdfplumber)
        
        pdfplumber_time = time.time() - start_time
        print(f"✅ pdfplumber 转换完成，耗时: {pdfplumber_time:.2f} 秒")
        
    except Exception as e:
        print(f"❌ pdfplumber 转换失败: {e}")
        pdfplumber_time = 0
    
    # 测试PyMuPDF (如果可用)
    print("\n" + "="*50)
    print("🔧 测试 PyMuPDF")
    print("="*50)
    
    try:
        import fitz
        
        output_pymupdf = "/tmp/test_pymupdf_output.docx"
        start_time = time.time()
        
        doc = fitz.open(pdf_path)
        word_doc = Document()
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                paragraph = word_doc.add_paragraph(text)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        word_doc.save(output_pymupdf)
        doc.close()
        
        pymupdf_time = time.time() - start_time
        print(f"✅ PyMuPDF 转换完成，耗时: {pymupdf_time:.2f} 秒")
        
    except Exception as e:
        print(f"❌ PyMuPDF 转换失败: {e}")
        pymupdf_time = 0
    
    # 测试pdf2docx (如果可用)
    print("\n" + "="*50)
    print("📦 测试 pdf2docx")
    print("="*50)
    
    try:
        from pdf2docx import Converter
        
        output_pdf2docx = "/tmp/test_pdf2docx_output.docx"
        start_time = time.time()
        
        cv = Converter(pdf_path)
        cv.convert(output_pdf2docx)
        cv.close()
        
        pdf2docx_time = time.time() - start_time
        print(f"✅ pdf2docx 转换完成，耗时: {pdf2docx_time:.2f} 秒")
        
    except Exception as e:
        print(f"❌ pdf2docx 转换失败: {e}")
        pdf2docx_time = 0
    
    # 对比结果
    print("\n" + "="*50)
    print("📊 对比结果")
    print("="*50)
    print(f"pdfminer (格式化) 耗时: {pdfminer_time1:.2f} 秒")
    print(f"pdfminer (简单) 耗时: {pdfminer_time2:.2f} 秒")
    print(f"pdfplumber 耗时: {pdfplumber_time:.2f} 秒")
    print(f"PyMuPDF 耗时: {pymupdf_time:.2f} 秒")
    print(f"pdf2docx 耗时: {pdf2docx_time:.2f} 秒")
    
    # 性能分析
    times = [
        ("pdfminer (格式化)", pdfminer_time1),
        ("pdfminer (简单)", pdfminer_time2),
        ("pdfplumber", pdfplumber_time),
        ("PyMuPDF", pymupdf_time),
        ("pdf2docx", pdf2docx_time)
    ]
    
    # 找到最快的方法
    valid_times = [(name, time) for name, time in times if time > 0]
    if valid_times:
        fastest = min(valid_times, key=lambda x: x[1])
        print(f"\n🏆 最快方法: {fastest[0]} ({fastest[1]:.2f}秒)")
        
        # 计算相对性能
        for name, time_val in valid_times:
            if time_val > 0 and name != fastest[0]:
                slowdown = ((time_val - fastest[1]) / fastest[1]) * 100
                print(f"   {name} 比最快方法慢 {slowdown:.1f}%")
    
    # 显示文件大小对比
    print("\n📊 文件大小对比:")
    for name, path in [
        ("pdfminer (格式化)", output_pdfminer_formatted),
        ("pdfminer (简单)", output_pdfminer_simple),
        ("pdfplumber", output_pdfplumber),
        ("PyMuPDF", output_pymupdf),
        ("pdf2docx", output_pdf2docx)
    ]:
        if os.path.exists(path):
            size = os.path.getsize(path)
            print(f"  {name}: {size/1024:.2f} KB")

def main():
    """主函数"""
    print("🧪 pdfminer PDF转DOCX测试脚本")
    print("="*50)
    
    # 检查命令行参数
    if len(sys.argv) < 2:
        print("❌ 用法: python test_pdfminer_conversion.py <PDF文件路径>")
        print("📝 示例: python test_pdfminer_conversion.py /path/to/test.pdf")
        sys.exit(1)
    
    pdf_path = sys.argv[1]
    
    if not os.path.exists(pdf_path):
        print(f"❌ PDF文件不存在: {pdf_path}")
        sys.exit(1)
    
    print(f"📁 测试文件: {pdf_path}")
    print(f"📊 文件大小: {os.path.getsize(pdf_path) / 1024:.2f} KB")
    
    # 1. 测试PDF文本提取
    print("\n" + "="*50)
    print("🔍 步骤1: PDF文本提取测试")
    print("="*50)
    has_text, text_content, page_lengths = test_pdfminer_text_extraction(pdf_path)
    
    # 2. 测试PDF布局分析
    print("\n" + "="*50)
    print("📐 步骤2: PDF布局分析测试")
    print("="*50)
    text_boxes, figures, images = test_pdfminer_layout_analysis(pdf_path)
    
    # 3. 判断PDF类型
    print("\n" + "="*50)
    print("📋 步骤3: PDF类型判断")
    print("="*50)
    if has_text:
        print("✅ 文本版PDF - 包含可提取文本")
    else:
        print("📷 扫描版PDF - 主要包含图片")
    
    print(f"📊 文本框数量: {text_boxes}")
    print(f"📊 图形数量: {figures}")
    print(f"🖼️ 图片数量: {images}")
    
    # 4. 执行转换测试
    print("\n" + "="*50)
    print("🚀 步骤4: 转换测试")
    print("="*50)
    test_pdfminer_vs_others(pdf_path)
    
    print("\n" + "="*50)
    print("✅ 测试完成!")
    print("="*50)

if __name__ == "__main__":
    main()
