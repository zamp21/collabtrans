#!/usr/bin/env python3
"""
pdfplumber PDF to DOCX conversion test script
使用pdfplumber进行PDF转DOCX转换测试
"""

import pdfplumber
import time
import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
from PIL import Image

def test_pdfplumber_text_extraction(pdf_path):
    """测试pdfplumber文本提取"""
    print(f"🔍 测试pdfplumber文本提取: {pdf_path}")
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            total_text = ""
            page_count = len(pdf.pages)
            
            print(f"📄 PDF页数: {page_count}")
            
            for page_num in range(min(3, page_count)):  # 只检查前3页
                page = pdf.pages[page_num]
                text = page.extract_text()
                if text:
                    total_text += text
                    print(f"   第{page_num + 1}页文本长度: {len(text)} 字符")
                else:
                    print(f"   第{page_num + 1}页: 无文本内容")
            
            print(f"📊 总文本长度: {len(total_text)} 字符")
            print(f"📝 前100字符预览: {total_text[:100]}...")
            
            return len(total_text.strip()) > 50, total_text
            
    except Exception as e:
        print(f"❌ 文本提取失败: {e}")
        return False, ""

def test_pdfplumber_table_extraction(pdf_path):
    """测试pdfplumber表格提取"""
    print(f"📊 测试pdfplumber表格提取: {pdf_path}")
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            total_tables = 0
            
            for page_num in range(min(3, len(pdf.pages))):  # 只检查前3页
                page = pdf.pages[page_num]
                tables = page.extract_tables()
                total_tables += len(tables)
                print(f"   第{page_num + 1}页表格数量: {len(tables)}")
                
                # 显示第一个表格的结构
                if tables:
                    first_table = tables[0]
                    print(f"   第{page_num + 1}页第一个表格: {len(first_table)}行 x {len(first_table[0]) if first_table else 0}列")
            
            print(f"📊 总表格数量: {total_tables}")
            return total_tables
            
    except Exception as e:
        print(f"❌ 表格提取失败: {e}")
        return 0

def test_pdfplumber_image_extraction(pdf_path):
    """测试pdfplumber图片提取"""
    print(f"🖼️ 测试pdfplumber图片提取: {pdf_path}")
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            total_images = 0
            
            for page_num in range(min(3, len(pdf.pages))):  # 只检查前3页
                page = pdf.pages[page_num]
                images = page.images
                total_images += len(images)
                print(f"   第{page_num + 1}页图片数量: {len(images)}")
                
                # 显示图片信息
                for img_index, img in enumerate(images):
                    print(f"     图片{img_index + 1}: {img['width']}x{img['height']} 位置: ({img['x0']:.1f}, {img['y0']:.1f})")
            
            print(f"📊 总图片数量: {total_images}")
            return total_images
            
    except Exception as e:
        print(f"❌ 图片提取失败: {e}")
        return 0

def convert_pdf_to_docx_pdfplumber(pdf_path, output_path, preserve_tables=True, preserve_images=True):
    """使用pdfplumber转换PDF到DOCX"""
    print(f"🚀 开始pdfplumber转换: {pdf_path} -> {output_path}")
    
    start_time = time.time()
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            page_count = len(pdf.pages)
            print(f"📄 处理 {page_count} 页")
            
            # 创建Word文档
            word_doc = Document()
            
            processed_pages = 0
            total_tables = 0
            total_images = 0
            total_text_length = 0
            
            for page_num in range(page_count):
                page = pdf.pages[page_num]
                print(f"   处理第 {page_num + 1} 页...")
                
                # 提取文本
                text = page.extract_text()
                if text and text.strip():
                    total_text_length += len(text)
                    
                    # 添加文本到Word文档
                    lines = text.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:
                            paragraph = word_doc.add_paragraph(line)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # 处理表格
                if preserve_tables:
                    tables = page.extract_tables()
                    for table_data in tables:
                        if table_data and len(table_data) > 0:
                            # 创建Word表格
                            word_table = word_doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                            word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                            
                            # 填充表格数据
                            for row_idx, row_data in enumerate(table_data):
                                for col_idx, cell_data in enumerate(row_data):
                                    if cell_data:
                                        word_table.cell(row_idx, col_idx).text = str(cell_data).strip()
                            
                            total_tables += 1
                            
                            # 在表格后添加空行
                            word_doc.add_paragraph()
                
                # 处理图片
                if preserve_images:
                    images = page.images
                    for img_index, img in enumerate(images):
                        try:
                            # 获取图片区域
                            bbox = (img['x0'], img['y0'], img['x1'], img['y1'])
                            
                            # 裁剪页面为图片
                            cropped_page = page.crop(bbox)
                            
                            # 转换为PIL Image
                            pil_image = cropped_page.to_image(resolution=150)  # 150 DPI
                            
                            # 保存临时图片
                            temp_img_path = f"/tmp/temp_img_{page_num}_{img_index}_{int(time.time())}.png"
                            pil_image.save(temp_img_path)
                            
                            # 添加到Word文档
                            word_doc.add_picture(temp_img_path, width=Inches(4))
                            total_images += 1
                            
                            # 清理临时文件
                            os.remove(temp_img_path)
                            
                        except Exception as e:
                            print(f"   图片 {img_index} 处理失败: {e}")
                            continue
                
                # 添加分页符（除了最后一页）
                if page_num < page_count - 1:
                    word_doc.add_page_break()
                
                processed_pages += 1
            
            # 保存Word文档
            word_doc.save(output_path)
            
            end_time = time.time()
            duration = end_time - start_time
            
            print(f"✅ 转换完成! 耗时: {duration:.2f} 秒")
            print(f"📁 输出文件: {output_path}")
            print(f"📊 处理统计: {processed_pages}页, {total_tables}表格, {total_images}图片, {total_text_length}字符")
            
            # 检查输出文件大小
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                print(f"📊 输出文件大小: {file_size / 1024:.2f} KB")
            
            return True, duration, {
                'pages': processed_pages,
                'tables': total_tables,
                'images': total_images,
                'text_length': total_text_length,
                'file_size': file_size if os.path.exists(output_path) else 0
            }
            
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        return False, 0, {}

def test_pdfplumber_vs_others(pdf_path):
    """对比pdfplumber和其他转换方法"""
    print(f"🔄 对比测试: {pdf_path}")
    
    # 测试pdfplumber
    print("\n" + "="*50)
    print("📦 测试 pdfplumber")
    print("="*50)
    
    output_pdfplumber = "/tmp/test_pdfplumber_output.docx"
    success, pdfplumber_time, stats = convert_pdf_to_docx_pdfplumber(pdf_path, output_pdfplumber)
    
    if not success:
        pdfplumber_time = 0
        stats = {}
    
    # 测试pdf2docx (如果可用)
    print("\n" + "="*50)
    print("📦 测试 pdf2docx")
    print("="*50)
    
    try:
        from pdf2docx import Converter
        
        output_pdf2docx = "/tmp/test_pdf2docx_output.docx"
        start_time = time.time()
        
        cv = Converter(pdf_path)
        cv.convert(output_pdf2docx)
        cv.close()
        
        pdf2docx_time = time.time() - start_time
        print(f"✅ pdf2docx 转换完成，耗时: {pdf2docx_time:.2f} 秒")
        
    except Exception as e:
        print(f"❌ pdf2docx 转换失败: {e}")
        pdf2docx_time = 0
    
    # 测试PyMuPDF (如果可用)
    print("\n" + "="*50)
    print("🔧 测试 PyMuPDF")
    print("="*50)
    
    try:
        import fitz
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from PIL import Image
        import io
        
        output_pymupdf = "/tmp/test_pymupdf_output.docx"
        start_time = time.time()
        
        # 简单的PyMuPDF转换
        doc = fitz.open(pdf_path)
        word_doc = Document()
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                paragraph = word_doc.add_paragraph(text)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        word_doc.save(output_pymupdf)
        doc.close()
        
        pymupdf_time = time.time() - start_time
        print(f"✅ PyMuPDF 转换完成，耗时: {pymupdf_time:.2f} 秒")
        
    except Exception as e:
        print(f"❌ PyMuPDF 转换失败: {e}")
        pymupdf_time = 0
    
    # 对比结果
    print("\n" + "="*50)
    print("📊 对比结果")
    print("="*50)
    print(f"pdfplumber 耗时: {pdfplumber_time:.2f} 秒")
    print(f"pdf2docx 耗时: {pdf2docx_time:.2f} 秒")
    print(f"PyMuPDF 耗时: {pymupdf_time:.2f} 秒")
    
    if pdfplumber_time > 0:
        if pdf2docx_time > 0 and pdfplumber_time < pdf2docx_time:
            improvement = ((pdf2docx_time - pdfplumber_time) / pdf2docx_time) * 100
            print(f"🚀 pdfplumber 比 pdf2docx 快 {improvement:.1f}%")
        elif pdf2docx_time > 0:
            slowdown = ((pdfplumber_time - pdf2docx_time) / pdf2docx_time) * 100
            print(f"⚠️ pdfplumber 比 pdf2docx 慢 {slowdown:.1f}%")
        
        if pymupdf_time > 0 and pdfplumber_time < pymupdf_time:
            improvement = ((pymupdf_time - pdfplumber_time) / pymupdf_time) * 100
            print(f"🚀 pdfplumber 比 PyMuPDF 快 {improvement:.1f}%")
        elif pymupdf_time > 0:
            slowdown = ((pdfplumber_time - pymupdf_time) / pymupdf_time) * 100
            print(f"⚠️ pdfplumber 比 PyMuPDF 慢 {slowdown:.1f}%")
    
    # 显示文件大小对比
    print("\n📊 文件大小对比:")
    for name, path in [("pdfplumber", output_pdfplumber), ("pdf2docx", output_pdf2docx), ("PyMuPDF", output_pymupdf)]:
        if os.path.exists(path):
            size = os.path.getsize(path)
            print(f"  {name}: {size/1024:.2f} KB")

def main():
    """主函数"""
    print("🧪 pdfplumber PDF转DOCX测试脚本")
    print("="*50)
    
    # 检查命令行参数
    if len(sys.argv) < 2:
        print("❌ 用法: python test_pdfplumber_conversion.py <PDF文件路径>")
        print("📝 示例: python test_pdfplumber_conversion.py /path/to/test.pdf")
        sys.exit(1)
    
    pdf_path = sys.argv[1]
    
    if not os.path.exists(pdf_path):
        print(f"❌ PDF文件不存在: {pdf_path}")
        sys.exit(1)
    
    print(f"📁 测试文件: {pdf_path}")
    print(f"📊 文件大小: {os.path.getsize(pdf_path) / 1024:.2f} KB")
    
    # 1. 测试PDF文本提取
    print("\n" + "="*50)
    print("🔍 步骤1: PDF文本提取测试")
    print("="*50)
    has_text, text_content = test_pdfplumber_text_extraction(pdf_path)
    
    # 2. 测试PDF表格提取
    print("\n" + "="*50)
    print("📊 步骤2: PDF表格提取测试")
    print("="*50)
    table_count = test_pdfplumber_table_extraction(pdf_path)
    
    # 3. 测试PDF图片提取
    print("\n" + "="*50)
    print("🖼️ 步骤3: PDF图片提取测试")
    print("="*50)
    image_count = test_pdfplumber_image_extraction(pdf_path)
    
    # 4. 判断PDF类型
    print("\n" + "="*50)
    print("📋 步骤4: PDF类型判断")
    print("="*50)
    if has_text:
        print("✅ 文本版PDF - 包含可提取文本")
    else:
        print("📷 扫描版PDF - 主要包含图片")
    
    print(f"📊 表格数量: {table_count}")
    print(f"🖼️ 图片数量: {image_count}")
    
    # 5. 执行转换测试
    print("\n" + "="*50)
    print("🚀 步骤5: 转换测试")
    print("="*50)
    test_pdfplumber_vs_others(pdf_path)
    
    print("\n" + "="*50)
    print("✅ 测试完成!")
    print("="*50)

if __name__ == "__main__":
    main()
